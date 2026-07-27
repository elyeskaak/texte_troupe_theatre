"""
Génère les quatre notebooks Colab de `notebooks/`.

Pourquoi un générateur plutôt que des `.ipynb` écrits à la main : le format
Jupyter est un JSON verbeux où la moindre virgule manquante rend le fichier
illisible, et où le contenu utile est noyé dans les métadonnées. Décrire les
cellules en Python garantit un JSON valide, rend les quatre notebooks
rigoureusement homogènes, et permet de corriger le préambule des quatre en une
seule édition.

    python outils/generer_notebooks.py

À relancer après toute modification de ce fichier. Les notebooks produits sont
versionnés : c'est eux que l'on ouvre dans Colab.
"""

from __future__ import annotations

import json
from pathlib import Path

RACINE = Path(__file__).resolve().parent.parent
DOSSIER_NOTEBOOKS = RACINE / "notebooks"

DEPOT_COMPTE = "elyeskaak"
DEPOT_NOM = "texte_troupe_theatre"

# Le dépôt est privé : le clone depuis Colab exige un jeton d'accès personnel.
NOM_SECRET_JETON = "GITHUB_TOKEN"


# ============================================================
# FABRIQUES DE CELLULES
# ============================================================


def markdown(texte: str) -> dict:
    """Crée une cellule Markdown."""
    return {
        "cell_type": "markdown",
        "metadata": {},
        "source": texte.strip().split("\n"),
    }


def code(texte: str) -> dict:
    """Crée une cellule de code, non exécutée."""
    return {
        "cell_type": "code",
        "execution_count": None,
        "metadata": {},
        "outputs": [],
        "source": texte.strip().split("\n"),
    }


def notebook(titre: str, cellules: list[dict]) -> dict:
    """Assemble un notebook complet, avec les métadonnées attendues par Colab."""
    return {
        "nbformat": 4,
        "nbformat_minor": 0,
        "metadata": {
            "colab": {"name": titre, "provenance": [], "toc_visible": True},
            "kernelspec": {"display_name": "Python 3", "name": "python3"},
            "language_info": {"name": "python"},
        },
        "cells": cellules,
    }


# ============================================================
# CELLULES COMMUNES
# ============================================================


def preambule(numero: str, titre: str, description: str) -> list[dict]:
    """
    Cellules d'ouverture, identiques pour les quatre notebooks.

    Le notebook n'est qu'une interface : il monte le Drive, installe les
    dépendances, rend le paquet importable, puis appelle une fonction. Aucune
    logique métier n'y figure.
    """
    return [
        markdown(
            f"""
# {numero} — {titre}

{description}

---

**Ce notebook n'est qu'une interface.** Toute la logique vit dans le paquet
`theatre_editor`. On y monte le Drive, on installe les dépendances, on surcharge
éventuellement la configuration, puis on lance l'étape.

**Cette étape est reprenable.** Si Colab coupe, relancez la cellule
d'exécution : le travail déjà validé ne sera pas refait, et vous ne repaierez
aucun appel.
"""
        ),
        markdown("## 1. Dépendances et montage du Drive"),
        code(
            """
# Installation des dépendances du pipeline.
!pip install -q -U openai pymupdf python-docx

from google.colab import drive

drive.mount("/content/drive")
"""
        ),
        markdown(
            f"""
## 2. Récupération du code

Le dépôt [`{DEPOT_COMPTE}/{DEPOT_NOM}`](https://github.com/{DEPOT_COMPTE}/{DEPOT_NOM})
est **public** : rien à configurer, la cellule suivante suffit. Elle récupère la
dernière version du code à chaque exécution.

> **Si vous repassiez le dépôt en privé**, il faudrait un jeton d'accès :
> GitHub → *Settings* → *Developer settings* → *Personal access tokens* →
> *Fine-grained tokens*, avec **Contents : Read-only** sur ce seul dépôt. Puis
> l'enregistrer dans les Secrets de Colab sous le nom `{NOM_SECRET_JETON}`. La
> cellule le détecte et l'utilise automatiquement — aucune modification à faire.
"""
        ),
        code(
            f"""
# --- Option A : récupération depuis GitHub ------------------------------
DEPOT_COMPTE = "{DEPOT_COMPTE}"
DEPOT_NOM = "{DEPOT_NOM}"
DOSSIER_PROJET = f"/content/{{DEPOT_NOM}}"

import os
import subprocess
import sys

# Le jeton est OPTIONNEL : inutile sur un dépôt public, utilisé
# automatiquement s'il est présent. Ainsi la cellule fonctionne dans les deux
# cas, sans qu'il faille se souvenir de la visibilité du dépôt.
jeton = None

try:
    from google.colab import userdata

    jeton = userdata.get("{NOM_SECRET_JETON}")
except Exception:
    pass

# Quand un jeton est utilisé, l'URL le contient : elle ne doit JAMAIS être
# affichée ni figurer dans un message d'erreur. Les sorties de git sont donc
# capturées, jamais relayées.
if jeton:
    url = f"https://{{jeton}}@github.com/{{DEPOT_COMPTE}}/{{DEPOT_NOM}}.git"
else:
    url = f"https://github.com/{{DEPOT_COMPTE}}/{{DEPOT_NOM}}.git"

if os.path.isdir(DOSSIER_PROJET):
    commande = ["git", "-C", DOSSIER_PROJET, "pull", "--quiet"]
else:
    commande = ["git", "clone", "--quiet", url, DOSSIER_PROJET]

resultat = subprocess.run(commande, capture_output=True, text=True)

if resultat.returncode != 0:
    raise RuntimeError(
        "Récupération du code impossible.\\n"
        f"Vérifiez que le dépôt {{DEPOT_COMPTE}}/{{DEPOT_NOM}} est accessible.\\n"
        "S'il est privé, ajoutez un secret {NOM_SECRET_JETON} dans Colab."
    )

if DOSSIER_PROJET not in sys.path:
    sys.path.insert(0, DOSSIER_PROJET)

print("Code récupéré :", DOSSIER_PROJET)
print("Jeton GitHub  :", "utilisé" if jeton else "non nécessaire (dépôt public)")
"""
        ),
        code(
            """
# --- Option B : le dossier theatre_editor/ est sur votre Drive ---------
# Décommentez ces lignes et ajustez le chemin, puis n'exécutez PAS l'option A.

# import sys
# DOSSIER_PROJET = "/content/drive/MyDrive/texte_troupe_theatre"
# if DOSSIER_PROJET not in sys.path:
#     sys.path.insert(0, DOSSIER_PROJET)
"""
        ),
        markdown(
            """
## 3. Configuration

`config.py` porte toutes les valeurs par défaut. Les surcharges ci-dessous ne
valent que pour cette session : elles ne modifient pas le fichier.

**Vérifiez le dossier de travail** avant de continuer.
"""
        ),
        code(
            """
from pathlib import Path

from theatre_editor import config

# Dossier Drive contenant les PDF et recevant toutes les sorties.
config.DOSSIER_DRIVE = Path("/content/drive/MyDrive/Troupe 122 - 2026-27")

print("Dossier de travail :", config.DOSSIER_DRIVE)
print("Existe             :", config.DOSSIER_DRIVE.is_dir())
"""
        ),
    ]


CELLULE_CLE_API = code(
    """
# La clé API est lue depuis les Secrets de Colab.
#
#   panneau latéral « 🔑 Secrets » → ajouter OPENAI_API_KEY
#   → activer « Accès au notebook »
#
# Ainsi la clé n'apparaît jamais dans le notebook ni dans ses sorties.

from theatre_editor.utils import io

try:
    io.charger_cle_api()
    print("Clé API trouvée.")
except RuntimeError as erreur:
    print(erreur)
"""
)


CELLULE_VERIFIER_MODELES = code(
    """
# Contrôle que les identifiants de config.py existent bien sur ce compte.
# Deux secondes ici évitent de découvrir une faute de frappe après trois
# heures de traitement.

from theatre_editor.utils import api

api.verifier_modeles_configures()
"""
)


def cellule_journal(etape: str) -> dict:
    """Cellule d'inspection du journal d'une étape."""
    return code(
        f"""
# Journal détaillé de l'étape : un enregistrement par appel API, avec sa
# date, son modèle, son response_id, sa durée et sa consommation de jetons.

import json

chemin = config.DOSSIER_DRIVE / config.NOM_JOURNAL.format(etape="{etape}")

if chemin.exists():
    journal = json.loads(chemin.read_text(encoding="utf-8"))
    print("Dernière exécution :", journal["derniere_execution"])
    print("Configuration      :", json.dumps(journal["configuration"], ensure_ascii=False))
    print()

    for nom, bilan in journal["livres"].items():
        print(f"{{nom}} : {{json.dumps(bilan, ensure_ascii=False)}}")

    jetons = sum(
        (appel.get("tokens_entree") or 0) + (appel.get("tokens_sortie") or 0)
        for appel in journal["appels"]
    )
    print()
    print(f"{{len(journal['appels'])}} appel(s) journalisé(s), {{jetons:,}} jetons".replace(",", " "))
else:
    print("Aucun journal : l'étape n'a pas encore été lancée.")
"""
    )


# ============================================================
# NOTEBOOK 1 — OCR
# ============================================================


def notebook_ocr() -> dict:
    return notebook(
        "01_OCR.ipynb",
        [
            *preambule(
                "Étape 1",
                "OCR Vision",
                "Transcrit les PDF du dossier Drive en fichiers "
                "`<Livre>_OCR.txt`, page par page.\n\n"
                "**Le modèle ne corrige rien.** Il transcrit ce qu'il voit. "
                "C'est l'étape 2 qui corrigera, et c'est parce que `OCR.txt` "
                "reste brut qu'il pourra servir de référence à l'étape 3.",
            ),
            markdown("## 4. Clé API"),
            CELLULE_CLE_API,
            markdown("## 5. Vérification des modèles"),
            CELLULE_VERIFIER_MODELES,
            markdown(
                """
## 6. Aperçu du travail à faire

Liste les PDF trouvés et l'avancement de chacun, sans lancer aucun appel.
"""
            ),
            code(
                """
from theatre_editor.utils import io

for chemin in io.lister_pdf(config.DOSSIER_DRIVE):
    nom = io.nom_livre_depuis_pdf(chemin)
    chemins = io.resoudre_chemins(nom, chemin.parent)

    faites = sum(
        1
        for fichier in sorted(chemins.dossier_pages.glob("page_*.json"))
        if io.unite_terminee(fichier)
    ) if chemins.dossier_pages.is_dir() else 0

    taille_mo = chemin.stat().st_size / (1024 * 1024)
    print(f"{nom:<40} {taille_mo:>6.1f} Mo   {faites} page(s) déjà transcrite(s)")
"""
            ),
            markdown(
                """
## 7. Combien de pages seront réellement facturées ?

**Cette cellule ne consomme aucun jeton.**

Beaucoup de PDF ont déjà été passés à l'OCR par un scanner ou par Acrobat, et
portent donc une couche texte. Quand elle est de bonne qualité, le pipeline la
réutilise telle quelle : aucun appel API pour ces pages.

Mais une couche texte n'est pas forcément bonne — accents dépouillés, ligatures,
ordre de lecture faux. S'en servir à tort dégraderait tout le livre, puisque
l'étape 2 a pour consigne de ne pas réécrire l'auteur. Les contrôles sont donc
sévères, et le doute renvoie à l'OCR Vision.

Le diagnostic indique, pour chaque livre, combien de pages sont gratuites et
**pourquoi** les autres ne le sont pas.
"""
            ),
            code(
                """
from theatre_editor import ocr

diagnostics = ocr.diagnostiquer_couches_texte(config.DOSSIER_DRIVE)
"""
            ),
            markdown(
                """
### Ajuster la stratégie

`config.STRATEGIE_COUCHE_TEXTE` accepte trois valeurs.

| Valeur | Effet |
|---|---|
| `"auto"` | couche texte utilisée si elle passe les contrôles — **recommandé** |
| `"jamais"` | toujours l'OCR Vision, même sur un PDF déjà OCRisé |
| `"toujours"` | couche texte utilisée dès qu'elle existe, sans contrôle |

`"toujours"` est à réserver aux PDF dont vous connaissez la provenance et la
qualité. Sur un fichier douteux, il produirait un livre dégradé sans le signaler.

Si le diagnostic écarte beaucoup de pages pour un motif qui vous paraît trop
strict, vous pouvez relâcher le seuil correspondant — mais relisez alors un
extrait de la couche texte avant de lancer le livre entier.
"""
            ),
            code(
                """
config.STRATEGIE_COUCHE_TEXTE = "auto"

# Seuils de qualité, à ne relâcher qu'en connaissance de cause :
# config.MIN_CARACTERES_COUCHE_TEXTE = 200
# config.MIN_RATIO_ACCENTS = 0.005

print("Stratégie :", config.STRATEGIE_COUCHE_TEXTE)
"""
            ),
            markdown(
                """
### Inspecter une couche texte avant de lui faire confiance

Affiche ce que le PDF contient déjà pour une page donnée. À faire au moins une
fois sur un livre dont vous ne connaissez pas l'origine.
"""
            ),
            code(
                """
NOM_LIVRE = diagnostics[0].nom if diagnostics else None
NUMERO_PAGE = 5

if NOM_LIVRE:
    chemins = io.resoudre_chemins(NOM_LIVRE, config.DOSSIER_DRIVE)
    document = ocr.ouvrir_pdf(chemins.pdf)

    try:
        numero = min(NUMERO_PAGE, document.page_count)
        texte, raisons = ocr.evaluer_page_couche_texte(
            document.load_page(numero - 1)
        )
    finally:
        document.close()

    print(f"{NOM_LIVRE} — page {numero}")
    print("=" * 72)
    print(f"caractères extraits : {len(texte)}")
    print(f"retenue             : {ocr.couche_texte_retenue(texte, raisons)}")

    if raisons:
        print("motifs de refus     :")
        for raison in raisons:
            print("   -", raison)

    print("=" * 72)
    print(texte[:1200] if texte else "(aucune couche texte)")
"""
            ),
            markdown(
                """
## 8. Essai sur les premières pages

**À faire sur tout nouveau livre.** Éprouver les quatre étapes sur dix pages
coûte quelques centimes et révèle les mauvaises surprises — modèle qui refuse la
vision, couche texte trompeuse, structure mal reconnue — avant d'engager
trois cents pages.

Les pages transcrites pendant l'essai sont **conservées et réutilisées** lors du
passage complet : rien n'est perdu, rien n'est repayé.

Enchaînez ensuite les notebooks 02, 03 et 04 : ils travaillent depuis `OCR.txt`,
qui ne contiendra que les pages retenues. Aucun réglage à y reporter.
"""
            ),
            code(
                """
# Nombre de pages à traiter par PDF. None = livre entier.
config.LIMITE_PAGES = 10

print("Limite de pages :", config.LIMITE_PAGES)
"""
            ),
            markdown(
                """
### Passer au livre entier

Remettez `LIMITE_PAGES` à `None` et relancez la cellule de lancement. Seules les
pages manquantes seront transcrites.

Un point de vigilance, désormais géré par le code. Avec des blocs de 8 pages, un
essai de 10 pages produit un « bloc 2 » couvrant les pages 9 et 10, tandis que le
livre entier attend un bloc 2 couvrant les pages 9 à 16. L'étape 2 compare donc
les frontières enregistrées à celles recalculées, et **réédite tout bloc dont les
frontières ont changé** — sans quoi les pages 11 à 16 disparaîtraient
silencieusement.

Vous verrez alors s'afficher, ce qui est normal :

```
[ALERTE]  bloc 2 : frontières changées (pages 9–10 → 9–16), réédition
```
"""
            ),
            code(
                """
# Décommentez pour traiter le livre entier :

# config.LIMITE_PAGES = None
"""
            ),
            markdown(
                """
## 9. Lancement

Reprenable : relancez cette cellule autant de fois qu'il le faut.

Comptez environ 3 à 6 secondes par page. Un livre de 300 pages demande donc
entre 20 et 30 minutes, et la session Colab peut couper d'ici là — ce n'est pas
un problème.
"""
            ),
            code(
                """
from theatre_editor import ocr

resultats = ocr.executer(config.DOSSIER_DRIVE)
"""
            ),
            markdown(
                """
## 10. Contrôle du résultat

Affiche le début de chaque fichier produit, et signale les pages en échec.
"""
            ),
            code(
                """
for resultat in resultats:
    chemins = io.resoudre_chemins(resultat.nom, config.DOSSIER_DRIVE)
    print("=" * 72)
    print(resultat.nom, "—", resultat.statut)
    print("=" * 72)

    if resultat.numeros_echoues:
        print("Pages en échec :", resultat.numeros_echoues)
        print("Relancez la cellule 7 pour les reprendre.")
        print()

    if chemins.ocr.exists():
        print(io.lire_texte(chemins.ocr)[:1200])
"""
            ),
            markdown("## 11. Journal"),
            cellule_journal("ocr"),
        ],
    )


# ============================================================
# NOTEBOOK 2 — ÉDITION
# ============================================================


def notebook_edition() -> dict:
    return notebook(
        "02_Edition.ipynb",
        [
            *preambule(
                "Étape 2",
                "Édition OCR",
                "Transforme `<Livre>_OCR.txt` en `<Livre>_EDIT.txt` : "
                "correction des erreurs de reconnaissance, puis passe de "
                "raccord entre les blocs.\n\n"
                "**Le texte de l'auteur n'est jamais réécrit.** Seules les "
                "erreurs manifestes d'OCR sont corrigées.",
            ),
            markdown("## 4. Clé API"),
            CELLULE_CLE_API,
            markdown("## 5. Vérification des modèles"),
            CELLULE_VERIFIER_MODELES,
            markdown(
                """
## 6. Réglages de l'édition

`PAGES_PAR_BLOC` est le réglage le plus sensible. **Ne le changez pas au milieu
d'un livre** : les blocs déjà édités ne seraient plus alignés, et l'étape 3
refuserait de valider.
"""
            ),
            code(
                """
config.PAGES_PAR_BLOC = 8          # 6 à 10 est une bonne plage
config.LIGNES_CONTEXTE_RACCORD = 50
config.RATIO_MINIMAL_LONGUEUR = 0.80

print("Modèle d'édition :", config.MODEL_EDITION)
print("Modèle de raccord:", config.MODEL_RACCORD)
print("Pages par bloc   :", config.PAGES_PAR_BLOC)
"""
            ),
            markdown(
                """
## 7. Aperçu du découpage

Montre en combien de blocs chaque livre sera découpé, et ce qui est déjà fait.
"""
            ),
            code(
                """
from theatre_editor.utils import blocks, io

for chemin in io.lister_fichiers_ocr(config.DOSSIER_DRIVE):
    nom = io.nom_livre_depuis_ocr(chemin)
    chemins = io.resoudre_chemins(nom, chemin.parent)

    pages = blocks.decouper_en_pages(io.lire_texte(chemin))
    liste = blocks.former_blocs(pages, config.PAGES_PAR_BLOC)

    faits = sum(
        1 for b in liste if io.unite_terminee(chemins.bloc_json(b.numero))
    )
    raccords = sum(
        1
        for numero in range(1, max(1, len(liste)))
        if io.unite_terminee(chemins.raccord_json(numero))
    )

    print(f"{nom}")
    print(f"   {len(pages)} pages → {len(liste)} blocs")
    print(f"   {faits}/{len(liste)} bloc(s) édité(s), "
          f"{raccords}/{max(0, len(liste) - 1)} raccord(s) fait(s)")
"""
            ),
            markdown(
                """
## 8. Lancement

Les deux passes s'enchaînent : édition des blocs, puis raccord des jonctions.
Reprenable à l'unité près.
"""
            ),
            code(
                """
from theatre_editor import edition

resultats = edition.executer(config.DOSSIER_DRIVE)
"""
            ),
            markdown(
                """
## 9. Contrôle du résultat

Affiche le début de chaque `EDIT.txt` et la structure que l'étape 4 y verra.
C'est le moment de vérifier que la convention typographique est bien appliquée.
"""
            ),
            code(
                """
for resultat in resultats:
    chemins = io.resoudre_chemins(resultat.nom, config.DOSSIER_DRIVE)
    print("=" * 72)
    print(resultat.nom, "—", resultat.statut)
    print("=" * 72)

    if not chemins.edit.exists():
        print("Aucun fichier édité.")
        continue

    texte = io.lire_texte(chemins.edit)
    print(texte[:1200])
    print()
    print(blocks.rapport_classification(blocks.construire_index_structure(texte)))
"""
            ),
            markdown("## 10. Journal"),
            cellule_journal("edition"),
        ],
    )


# ============================================================
# NOTEBOOK 3 — VÉRIFICATION
# ============================================================


def notebook_verification() -> dict:
    return notebook(
        "03_Verification.ipynb",
        [
            *preambule(
                "Étape 3",
                "Contrôle qualité",
                "Compare `OCR.txt` et `EDIT.txt` pour détecter ce que "
                "l'édition aurait perdu, et produit `<Livre>_REPORT.txt`.\n\n"
                "**Le texte n'est jamais modifié.** Cette étape produit un "
                "diagnostic, pas une correction : c'est à vous de décider quoi "
                "faire de ce qu'elle signale.",
            ),
            markdown("## 4. Clé API"),
            CELLULE_CLE_API,
            markdown(
                """
## 5. Contrôles mécaniques d'abord

Ces contrôles sont **gratuits et instantanés** : aucun appel API. Lancez-les
seuls pour un premier avis, avant d'engager la comparaison par le modèle.
"""
            ),
            code(
                """
from theatre_editor.utils import blocks, io

for chemin in io.lister_fichiers_ocr(config.DOSSIER_DRIVE):
    nom = io.nom_livre_depuis_ocr(chemin)
    chemins = io.resoudre_chemins(nom, chemin.parent)

    if not chemins.edit.exists():
        print(f"{nom} : pas encore édité.")
        continue

    constats = blocks.controles_mecaniques(
        io.lire_texte(chemin), io.lire_texte(chemins.edit)
    )

    print(f"{nom} : {len(constats)} constat(s) mécanique(s)")
    for constat in constats:
        print("   ", constat)
"""
            ),
            markdown(
                """
## 6. Lancement de la comparaison complète

Un appel par bloc. Reprenable.
"""
            ),
            code(
                """
from theatre_editor import validation

resultats = validation.executer(config.DOSSIER_DRIVE)
"""
            ),
            markdown(
                """
## 7. Lecture du rapport

Le rapport est fait pour être lu par un humain. Il ne détaille que les blocs
porteurs de constats.
"""
            ),
            code(
                """
for resultat in resultats:
    chemins = io.resoudre_chemins(resultat.nom, config.DOSSIER_DRIVE)

    if chemins.report.exists():
        print(io.lire_texte(chemins.report))
    else:
        print(f"{resultat.nom} : aucun rapport produit.")
"""
            ),
            markdown(
                """
## 8. Que faire d'un constat ?

Le rapport signale, il ne corrige pas. Trois façons d'agir, de la plus légère à
la plus lourde.

**Corriger `EDIT.txt` à la main.** Le plus simple pour quelques constats
isolés. Le fichier est du texte, ouvrez-le et corrigez. Relancez ensuite
l'étape 4 seule.

**Refaire un bloc.** Supprimez ses fichiers dans `_EDIT_blocs/` et
`_EDIT_raccords/`, puis relancez l'étape 2 : seul ce bloc sera repayé.

**Réviser un prompt.** Si le même défaut revient sur beaucoup de blocs, c'est
le prompt qu'il faut corriger, dans `theatre_editor/prompts/`. Supprimez alors
`_EDIT_blocs/` en entier pour refaire le livre — et ne changez jamais de prompt
au milieu d'un livre, cela produirait des blocs hétérogènes.
"""
            ),
            code(
                """
# Exemple : refaire le bloc 12 d'un livre.
# Décommentez et ajustez le nom et le numéro.

# NOM_LIVRE = "Le Malentendu"
# NUMERO = 12
#
# chemins = io.resoudre_chemins(NOM_LIVRE, config.DOSSIER_DRIVE)
#
# for chemin in (
#     chemins.bloc_txt(NUMERO), chemins.bloc_json(NUMERO),
#     chemins.raccord_txt(NUMERO), chemins.raccord_json(NUMERO),
#     chemins.report_bloc_txt(NUMERO), chemins.report_bloc_json(NUMERO),
# ):
#     if chemin.exists():
#         chemin.unlink()
#         print("supprimé :", chemin.name)
"""
            ),
            markdown("## 9. Journal"),
            cellule_journal("validation"),
        ],
    )


# ============================================================
# NOTEBOOK 4 — DOCX
# ============================================================


def notebook_docx() -> dict:
    return notebook(
        "04_DOCX.ipynb",
        [
            *preambule(
                "Étape 4",
                "Génération DOCX",
                "Transforme `<Livre>_EDIT.txt` en `<Livre>.docx`.\n\n"
                "**Aucune IA, aucune clé API, aucun coût.** Uniquement "
                "`python-docx` et la convention typographique. Deux exécutions "
                "produisent le même document : régénérez autant que vous "
                "voulez après avoir ajusté un réglage.",
            ),
            markdown(
                """
## 4. Réglages typographiques

Modifiez librement : cette étape est gratuite et reproductible.
"""
            ),
            code(
                """
config.POLICE_TEXTE = "EB Garamond"
config.TAILLE_TEXTE_PT = 11
config.TAILLE_TITRE_ACTE_PT = 16
config.TAILLE_TITRE_SCENE_PT = 14
config.MARGE_CM = 3.0

config.SAUT_DE_PAGE_AVANT_ACTE = True
config.SAUT_DE_PAGE_AVANT_SCENE = False

for cle, definition in config.DEFINITIONS_STYLES.items():
    saut = " + saut de page" if definition["saut_de_page"] else ""
    graisse = "gras" if definition["gras"] else ("italique" if definition["italique"] else "romain")
    print(f"   {definition['nom']:<14} {definition['taille_pt']:>2} pt  "
          f"{definition['alignement']:<9} {graisse}{saut}")
"""
            ),
            markdown(
                """
## 5. Table d'inspection de la structure

**À lire avant de générer.** Elle montre comment chaque nom en gras a été
classé — acte, scène, personnage — et signale d'un `⚠` les classements
incertains.

C'est ici qu'on repère un acte pris pour un personnage, plutôt que de le
découvrir à la première page blanche parasite.
"""
            ),
            code(
                """
from theatre_editor.utils import blocks, io

for chemin in io.lister_fichiers_edit(config.DOSSIER_DRIVE):
    nom = io.nom_livre_depuis_edit(chemin)
    index = blocks.construire_index_structure(io.lire_texte(chemin))

    print("=" * 72)
    print(nom)
    print("=" * 72)
    print(blocks.rapport_classification(index))
    print()
"""
            ),
            markdown(
                """
## 6. Corriger un classement

Si la table ci-dessus se trompe, forcez le classement ici. Écrivez les noms
**en capitales et sans accents**, tels qu'ils apparaissent dans la colonne
`LABEL`.
"""
            ),
            code(
                """
# Exemples — décommentez et adaptez :

# config.PERSONNAGES_FORCES = frozenset({"LA VOIX", "LE CHOEUR"})
# config.TITRES_ACTE_FORCES = frozenset({"OUVERTURE"})
# config.TITRES_SCENE_FORCES = frozenset({"ENTRACTE"})

print("Personnages forcés :", sorted(config.PERSONNAGES_FORCES))
print("Actes forcés       :", sorted(config.TITRES_ACTE_FORCES))
print("Scènes forcées     :", sorted(config.TITRES_SCENE_FORCES))
"""
            ),
            markdown("## 7. Génération"),
            code(
                """
from theatre_editor import docx_export

resultats = docx_export.executer(config.DOSSIER_DRIVE)
"""
            ),
            markdown(
                """
## 8. Contrôle du document

Relit le DOCX produit et affiche le style appliqué à chaque paragraphe. Le
meilleur moyen de vérifier qu'actes, scènes et personnages sont bien distingués.
"""
            ),
            code(
                """
import docx

for resultat in resultats:
    chemins = io.resoudre_chemins(resultat.nom, config.DOSSIER_DRIVE)

    if not chemins.docx.exists():
        print(f"{resultat.nom} : aucun document produit.")
        continue

    document = docx.Document(str(chemins.docx))

    print("=" * 72)
    print(f"{resultat.nom} — {len(document.paragraphs)} paragraphes")
    print("=" * 72)

    for paragraphe in document.paragraphs[:40]:
        style = paragraphe.style.name.replace(config.PREFIXE_STYLE, "")
        saut = "  [PAGE NEUVE]" if paragraphe.style.paragraph_format.page_break_before else ""
        print(f"  {style:<14} | {paragraphe.text[:60]}{saut}")

    if len(document.paragraphs) > 40:
        print(f"  … {len(document.paragraphs) - 40} paragraphes de plus")
"""
            ),
            markdown(
                """
## 9. Téléchargement

Le document est déjà sur votre Drive. Cette cellule permet de le récupérer
directement sur votre machine.
"""
            ),
            code(
                """
from google.colab import files

for resultat in resultats:
    chemins = io.resoudre_chemins(resultat.nom, config.DOSSIER_DRIVE)

    if chemins.docx.exists():
        files.download(str(chemins.docx))
"""
            ),
            markdown(
                """
## 10. À propos de la police

`python-docx` inscrit le **nom** de la police dans le document, il ne
l'incorpore pas. EB Garamond n'a donc pas à être installée dans Colab pour que
la génération réussisse.

En revanche, si elle est absente de la machine qui **ouvre** le fichier, Word
substituera une autre police. Pour un rendu conforme, installez EB Garamond sur
votre poste — elle est gratuite et disponible sur Google Fonts.
"""
            ),
            markdown("## 11. Journal"),
            cellule_journal("docx"),
        ],
    )


# ============================================================
# ÉCRITURE
# ============================================================


NOTEBOOKS = {
    "01_OCR.ipynb": notebook_ocr,
    "02_Edition.ipynb": notebook_edition,
    "03_Verification.ipynb": notebook_verification,
    "04_DOCX.ipynb": notebook_docx,
}


def main() -> None:
    """Écrit les quatre notebooks."""
    DOSSIER_NOTEBOOKS.mkdir(parents=True, exist_ok=True)

    for nom, fabrique in NOTEBOOKS.items():
        chemin = DOSSIER_NOTEBOOKS / nom
        contenu = json.dumps(fabrique(), ensure_ascii=False, indent=1)

        # Fin de ligne LF explicite : le dépôt est normalisé, et un CRLF dans
        # un .ipynb produit des diffs illisibles.
        with open(chemin, "w", encoding="utf-8", newline="\n") as flux:
            flux.write(contenu + "\n")

        cellules = len(fabrique()["cells"])
        print(f"écrit : notebooks/{nom} ({cellules} cellules)")


if __name__ == "__main__":
    main()
