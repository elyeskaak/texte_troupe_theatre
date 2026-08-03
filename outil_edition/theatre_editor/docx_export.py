"""
Étape 4 — Génération DOCX : `<Livre>_EDIT.txt` → `<Livre>.docx`.

**Aucune IA.** Uniquement `python-docx` et la convention typographique établie
à l'étape 2. Deux exécutions sur le même `EDIT.txt` produisent le même
document : c'est ce déterminisme qui permet de régénérer un DOCX après avoir
changé une marge, sans repayer un seul appel.

Le traitement est en **deux temps**, et cet ordre n'est pas négociable :

1. `blocks.construire_index_structure()` parcourt le document entier et décide
   du type de chaque ligne en gras. Le type d'un `**UN.**` est une propriété
   *globale* — acte dans une pièce en parties numérotées, scène dans une pièce
   qui possède déjà des `ACTE I` — il ne peut donc pas se décider ligne à ligne.
2. Chaque ligne est ensuite classée et rendue selon son style.

Le parsing est lui-même à deux niveaux : le **niveau ligne** détermine le style
de paragraphe, le **niveau run** traite les emphases internes (`decouper_en_runs`).
Sans le second, une réplique du type « Je t'attendais *elle se lève* depuis une
heure » perdrait sa didascalie intercalée.

La table d'inspection est affichée **avant** la génération, afin que vous voyiez
ce que le parseur a compris — notamment les classements incertains — plutôt que
de le découvrir à la première page blanche parasite.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from functools import lru_cache
from pathlib import Path
from typing import Any

from theatre_editor import config
from theatre_editor.utils import blocks, io
from theatre_editor.utils import logging as journalisation

NOM_ETAPE = "docx"

# Types de ligne sans rendu propre : ils ne produisent aucun paragraphe.
TYPES_IGNORES = frozenset({blocks.TypeLigne.VIDE})

# Texte du séparateur de scène, rendu comme une didascalie centrée.
TEXTE_SEPARATEUR = "*"


# ============================================================
# 1. RÉSULTATS
# ============================================================


@dataclass
class ResultatLivre:
    """Bilan de la génération d'un document."""

    nom: str
    statut: str = config.STATUT_TERMINE
    paragraphes: int = 0
    actes: int = 0
    scenes: int = 0
    personnages: int = 0
    # Sortie de répétition. Restent à zéro si elle n'a pas pu être écrite.
    unites: int = 0
    repliques: int = 0
    classements_incertains: list[str] = field(default_factory=list)
    avertissements: list[str] = field(default_factory=list)
    duree_secondes: float = 0.0
    erreur: str | None = None

    def champs_journal(self) -> dict[str, Any]:
        return {
            "statut": self.statut,
            "paragraphes": self.paragraphes,
            "actes": self.actes,
            "scenes": self.scenes,
            "personnages": self.personnages,
            "unites": self.unites,
            "repliques": self.repliques,
            "classements_incertains": self.classements_incertains,
            "avertissements": self.avertissements,
            "duree_secondes": self.duree_secondes,
            "erreur": self.erreur,
        }


# ============================================================
# 2. DÉPENDANCE python-docx
# ============================================================


@lru_cache(maxsize=1)
def _module_docx():
    """
    Importe `python-docx` à la demande.

    Import différé pour la même raison que dans les autres modules : permettre
    de charger `theatre_editor` sans que toutes les dépendances soient
    installées.
    """
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.shared import Cm, Pt
        from docx.oxml.ns import qn

        return {
            "docx": docx,
            "ALIGN": WD_ALIGN_PARAGRAPH,
            "STYLE_TYPE": WD_STYLE_TYPE,
            "Cm": Cm,
            "Pt": Pt,
            "qn": qn,
        }
    except ImportError as erreur:  # pragma: no cover - dépend de l'environnement
        raise RuntimeError(
            "python-docx est introuvable.\n"
            "Installez-le :  pip install -U python-docx"
        ) from erreur


def _alignement(nom: str):
    """Traduit un alignement de `config` en énumération python-docx."""
    outils = _module_docx()

    correspondances = {
        "centre": outils["ALIGN"].CENTER,
        "justifie": outils["ALIGN"].JUSTIFY,
        "gauche": outils["ALIGN"].LEFT,
        # Aligné à droite : usage imprimé pour la source d'une épigraphe.
        "droite": outils["ALIGN"].RIGHT,
    }

    if nom not in correspondances:
        raise ValueError(
            f"alignement inconnu dans config.DEFINITIONS_STYLES : « {nom} ». "
            f"Valeurs admises : {', '.join(sorted(correspondances))}."
        )

    return correspondances[nom]


# ============================================================
# 3. CONSTRUCTION DU DOCUMENT
# ============================================================


def creer_document():
    """Crée un document vierge, avec ses marges et ses styles."""
    outils = _module_docx()

    document = outils["docx"].Document()

    _appliquer_marges(document)
    _creer_styles(document)

    return document


def _appliquer_marges(document) -> None:
    """
    Applique des marges généreuses sur les quatre côtés.

    Aucun numéro de page n'est ajouté : `python-docx` n'en insère aucun par
    défaut, l'exigence est donc satisfaite par abstention — il n'y a rien à
    faire, seulement rien à ajouter.
    """
    centimetres = _module_docx()["Cm"](config.MARGE_CM)

    for section in document.sections:
        section.top_margin = centimetres
        section.bottom_margin = centimetres
        section.left_margin = centimetres
        section.right_margin = centimetres


def _creer_styles(document) -> None:
    """
    Crée les styles de paragraphe décrits dans `config.DEFINITIONS_STYLES`.

    Tout passe par des styles nommés, jamais par une mise en forme appliquée
    run par run : changer le corps des titres de tout un livre reste ainsi la
    modification d'une seule ligne de `config.py`.
    """
    for cle, definition in config.DEFINITIONS_STYLES.items():
        _creer_style(document, cle, definition)


def _creer_style(document, cle: str, definition: dict[str, Any]) -> None:
    """Crée un style unique et le configure entièrement."""
    outils = _module_docx()

    nom = f"{config.PREFIXE_STYLE}{definition['nom']}"

    style = document.styles.add_style(nom, outils["STYLE_TYPE"].PARAGRAPH)
    style.quick_style = True

    police = style.font
    police.name = config.POLICE_TEXTE
    police.size = outils["Pt"](definition["taille_pt"])
    police.bold = bool(definition["gras"])
    police.italic = bool(definition["italique"])
    # Aucune couleur n'est définie : la valeur héritée est le noir automatique.

    _forcer_police(style, config.POLICE_TEXTE)

    format_paragraphe = style.paragraph_format
    format_paragraphe.alignment = _alignement(str(definition["alignement"]))
    format_paragraphe.space_before = outils["Pt"](definition["espace_avant_pt"])
    format_paragraphe.space_after = outils["Pt"](definition["espace_apres_pt"])
    format_paragraphe.page_break_before = bool(definition["saut_de_page"])

    # Évite qu'un nom de personnage se retrouve seul en bas de page, séparé de
    # la réplique qu'il annonce.
    format_paragraphe.keep_with_next = cle in {
        "titre_oeuvre",
        "titre_acte",
        "titre_scene",
        "titre_secondaire",
        "distribution",
        "personnage",
        "epigraphe",
    }


def _forcer_police(style, nom_police: str) -> None:
    """
    Renseigne les variantes `eastAsia` et `cs` de la police.

    `python-docx` ne remplit que l'attribut `ascii` de `w:rFonts`. Word peut
    alors substituer une autre police sur certains caractères — typiquement les
    guillemets ou les tirets cadratins d'un texte français. Ce détail est
    invisible à la génération et très visible à l'impression.
    """
    qualifier = _module_docx()["qn"]
    element = style.element.rPr.rFonts

    for attribut in ("w:eastAsia", "w:cs", "w:hAnsi"):
        element.set(qualifier(attribut), nom_police)


# ============================================================
# 4. RENDU DES LIGNES
# ============================================================


def nom_style(type_ligne: blocks.TypeLigne) -> str:
    """
    Retourne le nom du style associé à un type de ligne.

    Les valeurs de `TypeLigne` correspondant exactement aux clés de
    `DEFINITIONS_STYLES`, un type sans style provoque une erreur immédiate
    plutôt qu'un paragraphe au style par défaut passé inaperçu.
    """
    definition = config.DEFINITIONS_STYLES.get(type_ligne.value)

    if definition is None:
        raise KeyError(
            f"aucun style défini pour le type « {type_ligne.value} » : "
            "complétez config.DEFINITIONS_STYLES."
        )

    return f"{config.PREFIXE_STYLE}{definition['nom']}"


def ajouter_paragraphe(document, ligne: blocks.LigneClassee) -> bool:
    """
    Ajoute un paragraphe correspondant à une ligne classée.

    Returns:
        True si un paragraphe a été créé.
    """
    if ligne.type in TYPES_IGNORES:
        return False

    if ligne.type is blocks.TypeLigne.SEPARATEUR:
        paragraphe = document.add_paragraph(style=nom_style(blocks.TypeLigne.DIDASCALIE))
        run = paragraphe.add_run(TEXTE_SEPARATEUR)
        _appliquer_emphase_du_style(run, blocks.TypeLigne.DIDASCALIE)
        return True

    paragraphe = document.add_paragraph(style=nom_style(ligne.type))

    _ajouter_runs(paragraphe, ligne)

    return True


def _appliquer_emphase_du_style(run, type_ligne: blocks.TypeLigne) -> None:
    """
    Force le gras et l'italique du run d'après `DEFINITIONS_STYLES`, en plus
    de ce que porte déjà le style de paragraphe.

    Redondant avec le style **en théorie** : `python-docx` renseigne bien
    `style.font.italic`, et Word l'honore. Mais un défaut réel, constaté sur
    `le_dindon_feydeau_LT.docx`, a montré qu'un lecteur peut afficher une
    didascalie longue en romain alors que son style porte `italique=True` —
    tout en respectant correctement l'alignement du même style. La mise en
    forme au niveau du run, elle, ne dépend d'aucune interprétation du style
    par le lecteur.
    """
    definition = config.DEFINITIONS_STYLES[type_ligne.value]
    run.bold = bool(definition["gras"])
    run.italic = bool(definition["italique"])


def _ajouter_runs(paragraphe, ligne: blocks.LigneClassee) -> None:
    """
    Ajoute le contenu d'une ligne, en respectant ses emphases internes.

    Les emphases ne sont découpées que dans le corps du texte. Sur un titre,
    un nom de personnage ou une didascalie, le paragraphe est entièrement dans
    un seul registre : gras et italique s'appliquent donc au run entier,
    d'après la même définition que le style (voir `_appliquer_emphase_du_style`).
    """
    if ligne.type is not blocks.TypeLigne.TEXTE:
        run = paragraphe.add_run(ligne.texte)
        _appliquer_emphase_du_style(run, ligne.type)
        return

    for fragment in blocks.decouper_en_runs(ligne.texte):
        run = paragraphe.add_run(fragment.texte)

        # On ne force l'attribut que lorsqu'il est demandé : laisser None
        # préserve la valeur héritée du style.
        if fragment.gras:
            run.bold = True

        if fragment.italique:
            run.italic = True


# ============================================================
# 5. GÉNÉRATION D'UN LIVRE
# ============================================================


def construire_docx(
    texte: str,
    roles_liminaires: dict[int, blocks.TypeLigne] | None = None,
) -> tuple[Any, blocks.IndexStructure, int]:
    """
    Construit le document à partir d'un texte édité.

    Args:
        texte: contenu de `EDIT.txt`.
        roles_liminaires: rôles des premières lignes, produits par l'étape 2 bis.
            Absents, les liminaires retombent sur les règles déterministes —
            cette étape n'a donc **aucune dépendance** envers la précédente.

    Returns:
        `(document, index de structure, nombre de paragraphes)`.
    """
    index = blocks.construire_index_structure(texte)
    document = creer_document()

    lignes = lignes_classees(texte, index, roles_liminaires)

    paragraphes = 0

    for ligne in lignes:
        if ajouter_paragraphe(document, ligne):
            paragraphes += 1

    return document, index, paragraphes


def lignes_classees(
    texte: str,
    index: blocks.IndexStructure,
    roles_liminaires: dict[int, blocks.TypeLigne] | None = None,
) -> list[blocks.LigneClassee]:
    """
    Classe le document, rôles liminaires appliqués s'il y en a.

    Extrait de `construire_docx()` pour que la sortie de répétition parte
    **exactement** des mêmes lignes que le DOCX. Deux séquences de classification
    parallèles finiraient par diverger, et le JSON présenterait alors une
    structure que le document imprimé ne montre pas.
    """
    lignes = blocks.classifier_document(texte, index)

    if roles_liminaires:
        lignes = _appliquer_roles_liminaires(lignes, texte, roles_liminaires)

    return lignes


def _appliquer_roles_liminaires(
    lignes: list[blocks.LigneClassee],
    texte: str,
    roles: dict[int, blocks.TypeLigne],
) -> list[blocks.LigneClassee]:
    """
    Substitue les rôles annotés aux classements déterministes des liminaires.

    La correspondance se fait par **contenu brut** et non par indice : la
    classification peut dédoubler une ligne portant un nom et sa réplique, si
    bien que les positions de sortie ne correspondent plus à celles de l'entrée.
    Aligner sur les indices produirait des décalages silencieux.
    """
    source = texte.split("\n")
    par_brut: dict[str, blocks.TypeLigne] = {}

    for numero, type_ligne in roles.items():
        if 0 <= numero < len(source):
            par_brut[source[numero]] = type_ligne

    resultat: list[blocks.LigneClassee] = []
    dejaappliques: set[str] = set()

    for ligne in lignes:
        type_annote = par_brut.get(ligne.brut)

        # Une ligne dédoublée produit deux entrées de même `brut` : le rôle
        # annoté ne s'applique qu'à la première, l'autre gardant son type.
        if type_annote is not None and ligne.brut not in dejaappliques:
            dejaappliques.add(ligne.brut)
            resultat.append(
                blocks.LigneClassee(
                    brut=ligne.brut,
                    texte=blocks.contenu_sans_marqueurs(ligne.brut, type_annote),
                    type=type_annote,
                )
            )
            continue

        resultat.append(ligne)

    return resultat


def _ecrire_repet(
    *,
    chemins: io.CheminsLivre,
    texte: str,
    roles: dict[int, blocks.TypeLigne] | None,
    index: blocks.IndexStructure,
    resultat: ResultatLivre,
) -> None:
    """
    Écrit la sortie destinée à l'outil de répétition.

    **Le DOCX est déjà enregistré quand cette fonction est appelée, et c'est
    délibéré.** La génération du document imprimé est la raison d'être de
    l'étape ; la sortie de répétition en est un bénéfice annexe. Un défaut dans
    cette seconde sortie ne doit donc jamais coûter la première, ni faire passer
    le livre en échec.

    L'exception est néanmoins **signalée** comme avertissement, jamais avalée :
    un JSON manquant en silence serait découvert sur le téléphone, un dimanche
    de filage.
    """
    from theatre_editor import repet_export

    try:
        lignes = lignes_classees(texte, index, roles)
        document = repet_export.ecrire_repet(chemins, lignes, index)
    except Exception as erreur:
        resultat.avertissements.append(
            f"la sortie de répétition n'a pas pu être écrite ({erreur}) — "
            "le DOCX, lui, est bien généré"
        )
        return

    totaux = repet_export.compter(document)
    resultat.unites = totaux["unites"]
    resultat.repliques = totaux["repliques"]

    # Les anomalies de structure relevées à l'assemblage — un texte sans
    # personnage annoncé, par exemple — remontent dans le rapport de l'étape.
    # Sans cela, elles ne vivraient que dans un JSON que personne ne relit.
    nouveaux = [
        avertissement
        for avertissement in document["avertissements"]
        if avertissement not in resultat.avertissements
        and avertissement not in index.avertissements
    ]
    resultat.avertissements.extend(nouveaux)


def traiter_livre(
    chemins: io.CheminsLivre,
    journal: journalisation.Journal,
) -> ResultatLivre:
    """Génère le DOCX d'un livre à partir de son fichier édité."""
    nom_livre = chemins.nom
    resultat = ResultatLivre(nom=nom_livre)

    journalisation.section(f"DOCX — {nom_livre}")

    with journalisation.Chrono() as chrono:
        try:
            _generer(chemins=chemins, resultat=resultat)
        except Exception as erreur:
            resultat.statut = config.STATUT_ECHEC
            resultat.erreur = str(erreur)
            journalisation.echec(f"{nom_livre} : {erreur}")

    resultat.duree_secondes = chrono.secondes

    journal.resumer_livre(nom_livre, **resultat.champs_journal())
    journal.sauvegarder()

    return resultat


def _generer(*, chemins: io.CheminsLivre, resultat: ResultatLivre) -> None:
    """Corps de la génération d'un livre."""
    texte = io.lire_texte(chemins.edit)

    if not texte.strip():
        raise ValueError(f"{chemins.edit.name} est vide")

    # Rôles des liminaires, s'ils ont été annotés. Absents, la génération se
    # déroule exactement comme avant : cette étape reste gratuite et
    # déterministe, sans dépendance envers l'étape 2 bis.
    from theatre_editor import liminaires

    roles = liminaires.charger_roles(chemins)

    if roles:
        journalisation.info(f"   {len(roles)} rôle(s) liminaire(s) appliqué(s)")

    document, index, paragraphes = construire_docx(texte, roles)

    # La table d'inspection est affichée AVANT l'écriture : vous voyez ce que le
    # parseur a compris au moment où le document peut encore être rejeté.
    journalisation.info("")
    journalisation.info(blocks.rapport_classification(index))
    journalisation.info("")

    document.save(str(chemins.docx))

    resultat.paragraphes = paragraphes
    resultat.actes = index.compter(blocks.TypeLigne.TITRE_ACTE)
    resultat.scenes = index.compter(blocks.TypeLigne.TITRE_SCENE)
    resultat.personnages = index.compter(blocks.TypeLigne.PERSONNAGE)
    resultat.classements_incertains = [c.affichage for c in index.incertains]
    resultat.avertissements = list(index.avertissements)

    if config.MARQUEUR_ECHEC_BLOC.split("{")[0] in texte:
        resultat.avertissements.append(
            "le texte édité contient un marqueur de bloc en échec : "
            "relancez l'étape « edition »"
        )

    # Après l'affectation de `resultat.avertissements`, jamais avant : cette
    # ligne remplace la liste, et y ajouter quoi que ce soit plus tôt serait
    # écrasé sans trace. Le DOCX, lui, est déjà enregistré (voir _ecrire_repet).
    _ecrire_repet(
        chemins=chemins, texte=texte, roles=roles, index=index, resultat=resultat
    )

    for avertissement in resultat.avertissements:
        journalisation.alerte(avertissement)

    if resultat.repliques:
        journalisation.succes(
            f"{chemins.repet.name} — {resultat.unites} unité(s), "
            f"{resultat.repliques} réplique(s)"
        )

    journalisation.succes(
        f"{chemins.docx.name} — {paragraphes} paragraphes, "
        f"{resultat.actes} acte(s), {resultat.scenes} scène(s), "
        f"{resultat.personnages} personnage(s)"
    )


# ============================================================
# 6. POINT D'ENTRÉE DE L'ÉTAPE
# ============================================================


def executer(dossier: Path | None = None) -> list[ResultatLivre]:
    """
    Génère les DOCX de tous les livres édités du dossier.

    Args:
        dossier: dossier à parcourir. `config.DOSSIER_DRIVE` par défaut.

    Returns:
        Un bilan par livre traité.
    """
    base = dossier if dossier is not None else config.DOSSIER_DRIVE

    journalisation.titre("Étape 4 — Génération DOCX")

    livres = io.lister_livres_avec(config.NOM_EDIT, base)
    journalisation.info(f"Dossier : {base}")
    journalisation.info(f"Livres édités trouvés : {len(livres)}")

    if not livres:
        journalisation.alerte(
            f"aucun « {config.NOM_EDIT} » dans {config.DOSSIER_TEMPORAIRE}/ — "
            "lancez d'abord l'étape « edition »"
        )
        return []

    journal = journalisation.Journal.charger_ou_creer(
        NOM_ETAPE,
        base,
        {
            "police": config.POLICE_TEXTE,
            "taille_texte_pt": config.TAILLE_TEXTE_PT,
            "taille_titre_acte_pt": config.TAILLE_TITRE_ACTE_PT,
            "taille_titre_scene_pt": config.TAILLE_TITRE_SCENE_PT,
            "marge_cm": config.MARGE_CM,
            "saut_de_page_avant_acte": config.SAUT_DE_PAGE_AVANT_ACTE,
        },
    )

    resultats = [traiter_livre(chemins, journal) for chemins in livres]

    _afficher_recapitulatif(resultats, journal)

    return resultats


def _afficher_recapitulatif(
    resultats: list[ResultatLivre],
    journal: journalisation.Journal,
) -> None:
    """Affiche le bilan global de l'étape."""
    journalisation.recapitulatif(
        {
            "Documents générés": sum(
                1 for r in resultats if r.statut == config.STATUT_TERMINE
            ),
            "Échecs": sum(1 for r in resultats if r.statut == config.STATUT_ECHEC),
            "Paragraphes": sum(r.paragraphes for r in resultats),
            "Actes": sum(r.actes for r in resultats),
            "Scènes": sum(r.scenes for r in resultats),
            "Personnages": sum(r.personnages for r in resultats),
            "Répliques (répétition)": sum(r.repliques for r in resultats),
            "Durée": journalisation.formater_duree(
                sum(r.duree_secondes for r in resultats)
            ),
            "Journal": journal.chemin.name,
        }
    )

    incertains = [r.nom for r in resultats if r.classements_incertains]

    if incertains:
        journalisation.info("")
        journalisation.alerte(
            "classements incertains à relire dans la table d'inspection : "
            f"{', '.join(incertains)}"
        )
