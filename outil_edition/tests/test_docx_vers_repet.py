"""
Tests de `outils/docx_vers_repet.py`.

Le point qui compte ici n'est pas repris ailleurs : cet outil régénère le
`REPET.json` d'un DOCX déjà mis en forme **sans jamais produire ni écraser de
`.docx`**. `docx_vers_edit.py` et `repet_export.py` sont déjà couverts chacun
de leur côté ; ces tests portent sur l'assemblage des deux et sur l'absence de
`.docx` en sortie, qui est la raison d'être de ce module.
"""

from __future__ import annotations

import json
import sys
import tempfile
import unittest
from pathlib import Path

RACINE = Path(__file__).resolve().parent.parent

if str(RACINE) not in sys.path:
    sys.path.insert(0, str(RACINE))

import docx  # noqa: E402

from outils.docx_vers_repet import main, regenerer_manifeste, regenerer_repet  # noqa: E402
from theatre_editor import config  # noqa: E402
from theatre_editor.utils import io  # noqa: E402


def _ecrire_docx_minimal(chemin: Path, *, nom: str = "JAN", texte: str = "Nous y sommes enfin.") -> None:
    """Un DOCX minimal, avec exactement ce que `docx_vers_edit` sait lire."""
    document = docx.Document()
    document.add_paragraph("ACTE PREMIER", style="Heading 1")

    personnage = document.add_paragraph()
    personnage.add_run(nom).bold = True

    document.add_paragraph(texte)

    document.save(str(chemin))


class SansDocxEnSortie(unittest.TestCase):
    """L'invariant qui justifie l'existence de ce module."""

    def test_aucun_docx_n_est_ecrit(self):
        with tempfile.TemporaryDirectory() as brut:
            dossier = Path(brut)
            source = dossier / "source" / "Ma pièce.docx"
            source.parent.mkdir(parents=True)
            _ecrire_docx_minimal(source)

            travail = dossier / "travail"
            regenerer_repet(source, travail)

            chemins = io.resoudre_chemins("Ma pièce", travail)

            self.assertFalse(chemins.docx.exists())
            self.assertTrue(chemins.repet.exists())
            self.assertTrue(chemins.edit.exists())

    def test_relancable_a_volonte(self):
        """Une deuxième exécution ne doit ni échouer, ni laisser de trace."""
        with tempfile.TemporaryDirectory() as brut:
            dossier = Path(brut)
            source = dossier / "source" / "Ma pièce.docx"
            source.parent.mkdir(parents=True)
            _ecrire_docx_minimal(source)

            travail = dossier / "travail"
            regenerer_repet(source, travail)
            regenerer_repet(source, travail)

            chemins = io.resoudre_chemins("Ma pièce", travail)

            self.assertFalse(chemins.docx.exists())
            self.assertTrue(chemins.repet.exists())


class ContenuDuRepet(unittest.TestCase):
    """Le JSON produit est celui qu'écrirait l'étape 4, à l'identique."""

    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        dossier = Path(self._tmp.name)
        self.addCleanup(self._tmp.cleanup)

        source = dossier / "source" / "Ma pièce.docx"
        source.parent.mkdir(parents=True)
        _ecrire_docx_minimal(source)

        self.travail = dossier / "travail"
        self.chemin_repet = regenerer_repet(source, self.travail)

    def test_le_schema_et_la_piece_sont_corrects(self):
        document = json.loads(self.chemin_repet.read_text(encoding="utf-8"))

        self.assertEqual(document["schema"], config.SCHEMA_REPET)
        self.assertEqual(document["piece"], "Ma pièce")

    def test_la_replique_est_presente(self):
        document = json.loads(self.chemin_repet.read_text(encoding="utf-8"))
        repliques = [
            e
            for u in document["unites"]
            for e in u["elements"]
            if e["type"] == "replique"
        ]

        self.assertEqual(len(repliques), 1)
        self.assertEqual(repliques[0]["personnages"], ["JAN"])
        self.assertEqual(repliques[0]["texte"], "Nous y sommes enfin.")


class Manifeste(unittest.TestCase):
    """
    `manifest.json` — la liste des pièces disponibles dans le dossier partagé,
    lue par `outil_repetition` et `outil_lecture` au démarrage.
    """

    def test_recense_les_pieces_presentes(self):
        with tempfile.TemporaryDirectory() as brut:
            dossier = Path(brut)
            source = dossier / "source" / "Ma pièce.docx"
            source.parent.mkdir(parents=True)
            _ecrire_docx_minimal(source)

            travail = dossier / "travail"
            regenerer_repet(source, travail)

            chemin_manifeste = regenerer_manifeste(travail)
            manifeste = json.loads(chemin_manifeste.read_text(encoding="utf-8"))

            self.assertEqual(len(manifeste["pieces"]), 1)
            entree = manifeste["pieces"][0]
            self.assertEqual(entree["piece"], "Ma pièce")
            self.assertEqual(entree["fichier"], "Ma pièce_REPET.json")
            self.assertEqual(entree["unites"], 1)
            self.assertEqual(entree["repliques"], 1)
            self.assertEqual(entree["personnages"], 1)
            self.assertIn("genere_le", manifeste)

    def test_recense_tout_le_dossier_pas_seulement_le_dernier_lot(self):
        """
        Le manifeste reflète l'état du dossier, pas l'historique des appels :
        une pièce convertie hier doit rester listée aujourd'hui, même si on ne
        relance la conversion que d'une autre pièce.
        """
        with tempfile.TemporaryDirectory() as brut:
            dossier = Path(brut)
            travail = dossier / "travail"

            source_a = dossier / "source" / "Piece A.docx"
            source_a.parent.mkdir(parents=True)
            _ecrire_docx_minimal(source_a)
            regenerer_repet(source_a, travail)
            regenerer_manifeste(travail)

            source_b = dossier / "source" / "Piece B.docx"
            _ecrire_docx_minimal(source_b)
            regenerer_repet(source_b, travail)

            chemin_manifeste = regenerer_manifeste(travail)
            manifeste = json.loads(chemin_manifeste.read_text(encoding="utf-8"))

            noms = {p["piece"] for p in manifeste["pieces"]}
            self.assertEqual(noms, {"Piece A", "Piece B"})

    def test_un_repet_illisible_est_ignore_sans_faire_echouer_le_reste(self):
        with tempfile.TemporaryDirectory() as brut:
            dossier = Path(brut)
            source = dossier / "source" / "Ma pièce.docx"
            source.parent.mkdir(parents=True)
            _ecrire_docx_minimal(source)

            travail = dossier / "travail"
            travail.mkdir()
            regenerer_repet(source, travail)

            casse = travail / "Cassee_REPET.json"
            casse.write_text("{ceci n'est pas du JSON", encoding="utf-8")

            chemin_manifeste = regenerer_manifeste(travail)
            manifeste = json.loads(chemin_manifeste.read_text(encoding="utf-8"))

            self.assertEqual(len(manifeste["pieces"]), 1)
            self.assertEqual(manifeste["pieces"][0]["piece"], "Ma pièce")


class MainEtManifeste(unittest.TestCase):
    """L'écriture du manifeste depuis le point d'entrée CLI."""

    def test_avec_dossier_le_manifeste_est_ecrit(self):
        with tempfile.TemporaryDirectory() as brut:
            dossier = Path(brut)
            source = dossier / "source" / "Ma pièce.docx"
            source.parent.mkdir(parents=True)
            _ecrire_docx_minimal(source)

            travail = dossier / "travail"
            code = main([str(source), "--dossier", str(travail)])

            self.assertEqual(code, 0)
            self.assertTrue((travail / "manifest.json").exists())

    def test_dossier_est_obligatoire(self):
        """
        `--dossier` n'a pas de valeur par défaut qui fonctionne : le dossier du
        DOCX collisionnerait toujours avec le garde-fou de `convertir_fichier`
        (§ son docstring). Autant le rendre obligatoire plutôt que de laisser
        échouer chaque appel avec un message qui n'explique pas pourquoi.
        """
        with tempfile.TemporaryDirectory() as brut:
            dossier = Path(brut)
            source = dossier / "Ma pièce.docx"
            _ecrire_docx_minimal(source)

            with self.assertRaises(SystemExit):
                main([str(source)])


if __name__ == "__main__":  # pragma: no cover
    unittest.main()
