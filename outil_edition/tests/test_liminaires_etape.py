"""
Tests de l'étape 2 bis — rôles des pages liminaires.

Cette étape confie à un modèle le seul arbitrage que les règles déterministes ne
savent pas rendre : distinguer un titre, un auteur, une épigraphe, sa source, une
note d'éditeur, une liste de rôles et un prologue.

Trois propriétés sont vérifiées en priorité, car ce sont elles qui justifient de
découper cette passe à part :

- **un appel par livre**, mis en cache, jamais un appel par bloc ;
- **l'étape 4 reste gratuite et déterministe** ;
- **la dégradation est propre** : sans annotation, le DOCX se génère comme avant.
"""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest import mock

from theatre_editor import config, docx_export, liminaires
from theatre_editor.utils import api, blocks, io

try:
    import docx  # noqa: F401

    DOCX_DISPONIBLE = True
except ImportError:  # pragma: no cover - dépend de l'environnement
    DOCX_DISPONIBLE = False


EDIT = (
    "**La mastication des morts**\n"
    "**Patrick Kermann**\n"
    "\n"
    "*Les morts ont le sommeil léger*\n"
    "Heiner Müller\n"
    "\n"
    "**PERSONNAGES**\n"
    "Gilles Rimey.\n"
    "Alphonsine Rouart.\n"
    "\n"
    "**JAN.**\n"
    "Bonjour.\n"
    "**JAN.**\n"
    "Encore.\n"
)

REPONSE = """
0|titre_oeuvre
1|titre_secondaire
3|epigraphe
4|attribution
6|distribution
7|entree_distribution
8|entree_distribution
"""


def resultat_api(texte: str) -> api.ResultatAppel:
    return api.ResultatAppel(
        texte=texte,
        modele=config.MODEL_LIMINAIRES,
        response_id="resp_test",
        tentative=1,
        duree_secondes=1.0,
        tokens_entree=400,
        tokens_sortie=60,
    )


class BaseLiminaires(unittest.TestCase):
    def setUp(self):
        self._verbosite = config.VERBOSITE
        config.VERBOSITE = 0

        self._dossier = tempfile.TemporaryDirectory()
        self.base = Path(self._dossier.name)
        self.chemins = io.resoudre_chemins("Kermann", self.base)

        io.ecrire_texte_atomique(self.chemins.edit, EDIT)

    def tearDown(self):
        self._dossier.cleanup()
        config.VERBOSITE = self._verbosite

    def executer(self, reponse: str = REPONSE):
        with mock.patch.object(
            api, "appeler_modele", side_effect=lambda **_kw: resultat_api(reponse)
        ) as appel:
            resultats = liminaires.executer(self.base)

        return resultats, appel


# ============================================================
# 1. ANALYSE DE LA RÉPONSE
# ============================================================


class TestInterpretation(unittest.TestCase):
    def test_roles_valides_retenus(self):
        roles, refuses = liminaires.interpreter_annotations(REPONSE)

        self.assertEqual(roles[0], "titre_oeuvre")
        self.assertEqual(roles[4], "attribution")
        self.assertEqual(refuses, [])

    def test_role_inconnu_refuse(self):
        """
        Un rôle inventé ne doit pas être propagé : il ne correspondrait à aucun
        style, et l'erreur n'apparaîtrait que plusieurs étapes plus loin.
        """
        roles, refuses = liminaires.interpreter_annotations(
            "0|titre_oeuvre\n1|colophon\n2|texte\n"
        )

        self.assertEqual(set(roles), {0, 2})
        self.assertEqual(refuses, ["colophon"])

    def test_bavardage_ignore(self):
        roles, _ = liminaires.interpreter_annotations(
            "Voici les rôles :\n0|titre_oeuvre\nJ'espère que cela convient.\n"
        )

        self.assertEqual(roles, {0: "titre_oeuvre"})

    def test_reponse_vide(self):
        self.assertEqual(liminaires.interpreter_annotations(""), ({}, []))


# ============================================================
# 2. FIN DES LIMINAIRES
# ============================================================


class TestFinDesLiminaires(unittest.TestCase):
    def test_s_arrete_au_premier_personnage_sur(self):
        index = blocks.construire_index_structure(EDIT)
        fin = blocks.fin_des_liminaires(EDIT, index)

        lignes = EDIT.split("\n")

        self.assertEqual(lignes[fin], "**JAN.**")

    def test_un_label_incertain_n_arrete_pas(self):
        """
        Les labels incertains sont précisément ceux que cette passe doit
        trancher : s'arrêter à eux reviendrait à ne rien lui soumettre.
        """
        index = blocks.construire_index_structure(EDIT)
        fin = blocks.fin_des_liminaires(EDIT, index)

        # « La mastication des morts » est incertain, en ligne 0.
        self.assertGreater(fin, 0)

    def test_plafonne_par_la_configuration(self):
        texte = "\n".join(f"Ligne {i}." for i in range(500))

        index = blocks.construire_index_structure(texte)

        self.assertLessEqual(
            blocks.fin_des_liminaires(texte, index), config.LIGNES_LIMINAIRES
        )

    def test_piece_sans_liminaire(self):
        texte = "**JAN.**\nBonjour.\n**JAN.**\nEncore.\n"

        index = blocks.construire_index_structure(texte)

        self.assertEqual(blocks.fin_des_liminaires(texte, index), 0)

    def test_titre_eclate_sur_plusieurs_lignes_grasses_n_arrete_pas(self):
        """
        Cas réel (Büchner, « La Mort de Danton ») : un titre imprimé sur
        plusieurs lignes grasses isolées et consécutives — « LA MORT. » /
        « DE. » / « DANTON. » — ne doit pas être pris pour le début de la
        pièce au seul motif que son dernier fragment coïncide avec le nom
        d'un vrai personnage. Sans la protection, le balayage s'arrêtait à
        « DANTON. » et les mentions d'éditeur qui suivaient devenaient sa
        première réplique.
        """
        texte = (
            "**LA MORT.**\n"
            "\n"
            "**DE.**\n"
            "\n"
            "**DANTON.**\n"
            "\n"
            "Mentions d'éditeur, traducteur, année.\n"
            "\n"
            "**PERSONNAGES**\n"
            "Danton. Julie.\n"
            "\n"
            "**DANTON.**\n"
            "Bonjour.\n"
            "**JULIE.**\n"
            "Bonjour.\n"
            "**DANTON.**\n"
            "Encore.\n"
        )

        index = blocks.construire_index_structure(texte)
        fin = blocks.fin_des_liminaires(texte, index)

        lignes = texte.split("\n")

        self.assertEqual(lignes[fin], "**DANTON.**")
        # La vraie première réplique, pas le fragment de titre (ligne 4).
        self.assertGreater(fin, 8)


# ============================================================
# 3. UN SEUL APPEL, MIS EN CACHE
# ============================================================


class TestCache(BaseLiminaires):
    def test_un_seul_appel_par_livre(self):
        _, appel = self.executer()

        self.assertEqual(appel.call_count, 1)

    def test_seconde_execution_n_appelle_plus_rien(self):
        self.executer()
        resultats, appel = self.executer()

        self.assertEqual(appel.call_count, 0)
        self.assertTrue(resultats[0].saute)

    def test_annotation_ecrite(self):
        self.executer()

        sidecar = io.lire_sidecar(self.chemins.liminaires)

        self.assertEqual(sidecar["roles"]["0"], "titre_oeuvre")
        self.assertGreater(sidecar["lignes_soumises"], 0)

    def test_piece_sans_liminaire_n_appelle_pas(self):
        io.ecrire_texte_atomique(
            self.chemins.edit, "**JAN.**\nBonjour.\n**JAN.**\nEncore.\n"
        )

        _, appel = self.executer()

        self.assertEqual(appel.call_count, 0)

    def test_echec_ne_bloque_pas_le_pipeline(self):
        """
        Cette étape est un raffinement : son échec ne doit pas empêcher de
        produire le DOCX, qui retombera sur les règles déterministes.
        """
        with mock.patch.object(
            api, "appeler_modele", side_effect=api.EchecAppelAPI("panne simulée")
        ):
            resultats = liminaires.executer(self.base)

        self.assertEqual(resultats[0].statut, config.STATUT_ECHEC)
        self.assertIsNone(io.lire_sidecar(self.chemins.liminaires))


# ============================================================
# 4. APPLICATION DANS LE DOCX
# ============================================================


@unittest.skipUnless(DOCX_DISPONIBLE, "python-docx n'est pas installé")
class TestApplicationDansLeDocx(BaseLiminaires):
    def _styles(self) -> list[tuple[str, str]]:
        document = docx.Document(str(self.chemins.docx))

        return [
            (p.style.name.replace(config.PREFIXE_STYLE, ""), p.text)
            for p in document.paragraphs
        ]

    def test_roles_appliques(self):
        self.executer()
        docx_export.executer(self.base)

        styles = dict((texte, style) for style, texte in self._styles())

        self.assertEqual(styles["La mastication des morts"], "Titre_Oeuvre")
        self.assertEqual(styles["Patrick Kermann"], "Titre_Secondaire")
        self.assertEqual(styles["Les morts ont le sommeil léger"], "Epigraphe")
        self.assertEqual(styles["Heiner Müller"], "Attribution")
        self.assertEqual(styles["PERSONNAGES"], "Distribution")

    def test_attribution_alignee_a_droite(self):
        """
        Le défaut que les règles déterministes ne savaient pas corriger : la
        source d'une épigraphe s'aligne à droite dans l'usage imprimé.
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self.executer()
        docx_export.executer(self.base)

        document = docx.Document(str(self.chemins.docx))
        attribution = next(
            p for p in document.paragraphs if p.text == "Heiner Müller"
        )

        self.assertEqual(
            attribution.style.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.RIGHT
        )

    def test_corps_de_la_piece_intact(self):
        """La passe ne doit toucher qu'aux liminaires."""
        self.executer()
        docx_export.executer(self.base)

        styles = dict((texte, style) for style, texte in self._styles())

        self.assertEqual(styles["JAN."], "Personnage")
        self.assertEqual(styles["Bonjour."], "Texte")

    def test_docx_genere_sans_annotation(self):
        """
        Dégradation propre : sans `LIMINAIRES.json`, la génération se déroule
        exactement comme avant. Cette étape est un raffinement, pas une
        dépendance.
        """
        docx_export.executer(self.base)

        self.assertTrue(self.chemins.docx.is_file())

        styles = dict((texte, style) for style, texte in self._styles())

        self.assertEqual(styles["JAN."], "Personnage")
        self.assertEqual(styles["La mastication des morts"], "Titre_Oeuvre")

    def test_generation_reste_deterministe(self):
        """
        L'annotation étant mise en cache, deux générations successives donnent
        le même document — sans nouvel appel.
        """
        self.executer()

        docx_export.executer(self.base)
        premiers = self._styles()

        with mock.patch.object(api, "appeler_modele") as appel:
            docx_export.executer(self.base)
            appel.assert_not_called()

        self.assertEqual(self._styles(), premiers)

    def test_aucune_asterisque_residuelle(self):
        self.executer()
        docx_export.executer(self.base)

        for _, texte in self._styles():
            if texte == docx_export.TEXTE_SEPARATEUR:
                continue

            with self.subTest(texte=texte):
                self.assertNotIn("*", texte)


if __name__ == "__main__":
    unittest.main(verbosity=2)
