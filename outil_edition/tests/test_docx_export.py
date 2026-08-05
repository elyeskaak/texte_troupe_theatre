"""
Tests de `theatre_editor.docx_export`.

Ces tests génèrent de **vrais fichiers DOCX** et les relisent avec
`python-docx`, plutôt que de vérifier des appels sur une doublure. C'est
possible parce que l'étape n'appelle aucune API, et c'est bien plus solide :
ce qui est contrôlé, c'est le contenu réellement écrit dans le document.

Le fil directeur est la conséquence visible de la classification : **un seul
saut de page, avant chaque acte**. Une scène prise pour un acte produirait une
page blanche parasite au milieu d'un acte, et c'est précisément ce que ces
tests rendent détectable.

Si `python-docx` est absent, le module entier est ignoré plutôt que d'échouer :
la suite doit rester exécutable sur une machine nue.
"""

from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from theatre_editor import config
from theatre_editor.utils import blocks, io

try:
    import docx  # noqa: F401

    from theatre_editor import docx_export

    DOCX_DISPONIBLE = True
except ImportError:  # pragma: no cover - dépend de l'environnement
    DOCX_DISPONIBLE = False


PIECE = (
    "**PERSONNAGES**\n"
    "JAN, le frère\n"
    "MARIA, sa femme\n"
    "LE MESSAGER\n"
    "\n"
    "\n"
    "**ACTE PREMIER**\n"
    "\n"
    "*Une auberge. Le soir.*\n"
    "\n"
    "**JAN.**\n"
    "Nous y sommes enfin.\n"
    "\n"
    "**MARIA.**\n"
    "Je t'attendais *elle se lève* depuis une heure.\n"
    "\n"
    "*Pause.*\n"
    "\n"
    "***\n"
    "\n"
    "**SCÈNE 2**\n"
    "\n"
    "**LE MESSAGER.**\n"
    "Un pli pour vous.\n"
    "\n"
    "**ACTE DEUXIÈME**\n"
    "\n"
    "**JAN.**\n"
    "Donnez.\n"
)


@unittest.skipUnless(DOCX_DISPONIBLE, "python-docx n'est pas installé")
class BaseDocx(unittest.TestCase):
    """Socle : génère un document et le relit."""

    TEXTE = PIECE

    def setUp(self):
        self._verbosite = config.VERBOSITE
        config.VERBOSITE = 0

        self._dossier = tempfile.TemporaryDirectory()
        self.base = Path(self._dossier.name)
        self.chemins = io.resoudre_chemins("Le Malentendu", self.base)

        io.ecrire_texte_atomique(self.chemins.edit, self.TEXTE)

        self.resultats = docx_export.executer(self.base)
        self.document = docx.Document(str(self.chemins.docx))

    def tearDown(self):
        self._dossier.cleanup()
        config.VERBOSITE = self._verbosite

    def paragraphes_du_style(self, cle: str) -> list:
        """Paragraphes portant le style identifié par sa clé de configuration."""
        nom = f"{config.PREFIXE_STYLE}{config.DEFINITIONS_STYLES[cle]['nom']}"

        return [p for p in self.document.paragraphs if p.style.name == nom]

    def textes_du_style(self, cle: str) -> list[str]:
        return [p.text for p in self.paragraphes_du_style(cle)]


# ============================================================
# 1. CLASSIFICATION RENDUE DANS LE DOCUMENT
# ============================================================


class TestClassificationRendue(BaseDocx):
    def test_actes_et_scenes_distingues(self):
        """
        Le nerf de l'affaire. Une confusion ici produit une page blanche
        parasite au milieu d'un acte.
        """
        self.assertEqual(
            self.textes_du_style("titre_acte"), ["ACTE PREMIER", "ACTE DEUXIÈME"]
        )
        self.assertEqual(self.textes_du_style("titre_scene"), ["SCÈNE 2"])

    def test_personnages_reconnus(self):
        self.assertEqual(
            sorted(set(self.textes_du_style("personnage"))),
            ["JAN.", "LE MESSAGER.", "MARIA."],
        )

    def test_entete_de_distribution_a_son_propre_style(self):
        self.assertEqual(self.textes_du_style("distribution"), ["PERSONNAGES"])

    def test_lieu_distingue_de_la_didascalie(self):
        """Un italique suivant un titre est un lieu ; ailleurs, une didascalie."""
        self.assertEqual(self.textes_du_style("lieu"), ["Une auberge. Le soir."])
        self.assertIn("Pause.", self.textes_du_style("didascalie"))

    def test_repliques_en_style_texte(self):
        textes = self.textes_du_style("texte")

        self.assertIn("Nous y sommes enfin.", textes)
        self.assertIn("Donnez.", textes)

    def test_aucune_asterisque_dans_le_document(self):
        """
        Les marqueurs de la convention ne doivent jamais atteindre le lecteur.
        Seul le séparateur de scène en conserve un, volontairement.
        """
        textes = [
            p.text
            for p in self.document.paragraphs
            if p.text != docx_export.TEXTE_SEPARATEUR
        ]

        for texte in textes:
            with self.subTest(texte=texte):
                self.assertNotIn("*", texte)

    def test_bilan_conforme(self):
        resultat = self.resultats[0]

        self.assertEqual(resultat.statut, config.STATUT_TERMINE)
        self.assertEqual(resultat.actes, 2)
        self.assertEqual(resultat.scenes, 1)
        self.assertEqual(resultat.personnages, 3)
        self.assertEqual(resultat.classements_incertains, [])


# ============================================================
# 2. SAUTS DE PAGE
# ============================================================


class TestSautsDePage(BaseDocx):
    """
    Le saut de page est porté par le **style**, non par le paragraphe.

    Conséquence à connaître pour lire ces tests : sur un paragraphe,
    `paragraph_format.page_break_before` vaut `None` — le formatage direct est
    correctement absent, la propriété étant héritée. C'est donc le style qu'il
    faut interroger.
    """

    def _style(self, cle: str):
        nom = f"{config.PREFIXE_STYLE}{config.DEFINITIONS_STYLES[cle]['nom']}"

        return self.document.styles[nom]

    def test_saut_de_page_sur_les_actes_seulement(self):
        for cle, definition in config.DEFINITIONS_STYLES.items():
            with self.subTest(style=cle):
                self.assertEqual(
                    bool(self._style(cle).paragraph_format.page_break_before),
                    bool(definition["saut_de_page"]),
                )

    def test_propriete_presente_dans_le_xml_du_style(self):
        """Garantie de plus bas niveau : c'est ce que Word lira réellement."""
        from docx.oxml.ns import qn

        pPr = self._style("titre_acte").element.find(qn("w:pPr"))

        self.assertIsNotNone(pPr)
        self.assertIsNotNone(pPr.find(qn("w:pageBreakBefore")))

    def test_les_scenes_ne_sautent_pas_de_page(self):
        """
        Vérification explicite du cas qui motive toute la classification à trois
        niveaux : une scène ne doit pas ouvrir une page neuve.
        """
        self.assertFalse(
            self._style("titre_scene").paragraph_format.page_break_before
        )

    def test_paragraphes_sans_formatage_direct(self):
        """
        Aucun saut n'est appliqué paragraphe par paragraphe : passer
        SAUT_DE_PAGE_AVANT_ACTE à False suffit donc à tout désactiver, sans
        résidu dans le document.
        """
        for paragraphe in self.document.paragraphs:
            with self.subTest(texte=paragraphe.text[:30]):
                self.assertIsNone(paragraphe.paragraph_format.page_break_before)

    def test_aucun_caractere_de_saut_de_page_insere(self):
        """Le saut ne doit pas être un `<w:br type="page"/>` glissé dans un run."""
        from docx.oxml.ns import qn

        for paragraphe in self.document.paragraphs:
            for run in paragraphe.runs:
                for saut in run._element.findall(qn("w:br")):
                    with self.subTest(texte=paragraphe.text[:30]):
                        self.assertNotEqual(saut.get(qn("w:type")), "page")


# ============================================================
# 3. EMPHASES INTERNES
# ============================================================


class TestEmphasesInternes(BaseDocx):
    def test_didascalie_intercalee_en_italique(self):
        """
        Sans le parsing au niveau des runs, cette réplique serait entièrement
        romaine et la didascalie perdue.
        """
        paragraphe = next(
            p for p in self.document.paragraphs if "attendais" in p.text
        )

        runs = [(r.text, bool(r.italic)) for r in paragraphe.runs]

        self.assertIn(("elle se lève", True), runs)
        self.assertTrue(any(t.startswith("Je t'attendais") and not i for t, i in runs))

    def test_texte_reconstitue_sans_perte(self):
        paragraphe = next(
            p for p in self.document.paragraphs if "attendais" in p.text
        )

        self.assertEqual(
            paragraphe.text, "Je t'attendais elle se lève depuis une heure."
        )


# ============================================================
# 4. STYLES ET MISE EN PAGE
# ============================================================


class TestStylesEtMiseEnPage(BaseDocx):
    def test_les_sept_styles_sont_crees(self):
        noms = {s.name for s in self.document.styles}

        for definition in config.DEFINITIONS_STYLES.values():
            attendu = f"{config.PREFIXE_STYLE}{definition['nom']}"
            with self.subTest(style=attendu):
                self.assertIn(attendu, noms)

    def test_hierarchie_des_corps(self):
        from docx.shared import Pt

        def taille(cle: str):
            nom = f"{config.PREFIXE_STYLE}{config.DEFINITIONS_STYLES[cle]['nom']}"
            return self.document.styles[nom].font.size

        self.assertEqual(taille("titre_acte"), Pt(20))
        self.assertEqual(taille("titre_scene"), Pt(18))
        self.assertEqual(taille("personnage"), Pt(15))
        self.assertEqual(taille("texte"), Pt(15))

    def test_police_appliquee_y_compris_ses_variantes(self):
        """
        `python-docx` ne renseigne que l'attribut `ascii`. Sans les variantes,
        Word substitue une autre police sur les guillemets et les tirets
        cadratins — invisible à la génération, très visible à l'impression.
        """
        from docx.oxml.ns import qn

        nom = f"{config.PREFIXE_STYLE}{config.DEFINITIONS_STYLES['texte']['nom']}"
        rfonts = self.document.styles[nom].element.rPr.rFonts

        for attribut in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            with self.subTest(attribut=attribut):
                self.assertEqual(rfonts.get(qn(attribut)), config.POLICE_TEXTE)

    def test_texte_justifie_et_titres_centres(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self.assertEqual(
            self.paragraphes_du_style("texte")[0].style.paragraph_format.alignment,
            WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        self.assertEqual(
            self.paragraphes_du_style("titre_acte")[0].style.paragraph_format.alignment,
            WD_ALIGN_PARAGRAPH.CENTER,
        )

    def test_marges_genereuses(self):
        """
        Comparaison au centième de centimètre, et non à l'identique.

        Le format DOCX stocke les marges en twips (1/20 de point) : 3 cm devient
        1701 twips, relu comme 1080135 EMU au lieu de 1080000. L'écart est de
        3,75 nanomètres et l'aller-retour exact est impossible — exiger
        l'égalité stricte serait exiger l'impossible.
        """
        section = self.document.sections[0]

        for nom, marge in (
            ("haut", section.top_margin),
            ("bas", section.bottom_margin),
            ("gauche", section.left_margin),
            ("droite", section.right_margin),
        ):
            with self.subTest(marge=nom):
                self.assertAlmostEqual(marge.cm, config.MARGE_CM, places=2)

    def test_aucune_couleur(self):
        """« Aucune couleur » : on ne définit jamais de couleur explicite."""
        for definition in config.DEFINITIONS_STYLES.values():
            nom = f"{config.PREFIXE_STYLE}{definition['nom']}"
            with self.subTest(style=nom):
                self.assertIsNone(self.document.styles[nom].font.color.rgb)

    def test_aucun_numero_de_page(self):
        """
        `python-docx` n'insère aucun numéro par défaut : l'exigence est
        satisfaite par abstention. Ce test garantit qu'aucun champ PAGE n'a été
        ajouté par inadvertance.
        """
        contenu = self.chemins.docx.read_bytes()

        self.assertNotIn(b"PAGE  \\* MERGEFORMAT", contenu)

    def test_nom_de_personnage_solidaire_de_sa_replique(self):
        """Un nom seul en bas de page, séparé de sa réplique, serait fautif."""
        for cle in ("titre_acte", "titre_scene", "personnage", "distribution"):
            with self.subTest(style=cle):
                self.assertTrue(
                    self.paragraphes_du_style(cle)[0].style.paragraph_format.keep_with_next
                )


# ============================================================
# 5. DÉTERMINISME
# ============================================================


class TestDeterminisme(BaseDocx):
    def test_deux_generations_donnent_le_meme_contenu(self):
        """
        Aucune IA : régénérer un DOCX après avoir changé une marge doit être
        strictement reproductible.
        """
        premiers = [(p.style.name, p.text) for p in self.document.paragraphs]

        docx_export.executer(self.base)
        seconds = [
            (p.style.name, p.text)
            for p in docx.Document(str(self.chemins.docx)).paragraphs
        ]

        self.assertEqual(premiers, seconds)

    def test_docx_ecrase_sans_accumulation(self):
        docx_export.executer(self.base)
        document = docx.Document(str(self.chemins.docx))

        self.assertEqual(
            len(document.paragraphs), len(self.document.paragraphs)
        )


# ============================================================
# 5 bis. SECONDE SORTIE : REPET.json
# ============================================================


class TestSortieDeRepetition(BaseDocx):
    """
    L'étape écrit aussi `<Livre>_REPET.json`, pour `../outil_repetition/`.

    Ce qui est vérifié ici, c'est le **câblage** : que la seconde sortie soit
    bien produite, et surtout qu'elle ne puisse jamais coûter la première. Le
    contenu du JSON, lui, est couvert par `test_repet_export.py`.
    """

    def test_le_json_est_ecrit_a_cote_du_docx(self):
        self.assertTrue(self.chemins.repet.exists())
        self.assertEqual(self.chemins.repet.parent, self.chemins.docx.parent)

    def test_le_json_decrit_la_meme_piece_que_le_docx(self):
        document = json.loads(self.chemins.repet.read_text(encoding="utf-8"))

        self.assertEqual(document["piece"], "Le Malentendu")
        self.assertEqual(document["schema"], config.SCHEMA_REPET)

        noms = {p["nom"] for p in document["personnages"]}

        self.assertEqual(noms, {"JAN", "MARIA", "LE MESSAGER"})

    def test_les_totaux_remontent_au_resultat(self):
        resultat = self.resultats[0]

        self.assertGreater(resultat.unites, 0)
        self.assertGreater(resultat.repliques, 0)

    def test_un_echec_du_json_ne_coute_pas_le_docx(self):
        """
        Le document imprimé est la raison d'être de l'étape ; la sortie de
        répétition en est un bénéfice annexe. L'échec de la seconde ne doit ni
        supprimer la première, ni faire passer le livre en échec — mais il doit
        être **signalé**, un JSON manquant en silence se découvrant sur le
        téléphone un dimanche de filage.
        """
        from theatre_editor import repet_export

        self.chemins.docx.unlink()
        self.chemins.repet.unlink()

        origine = repet_export.ecrire_repet

        def echouer(*_args, **_kwargs):
            raise RuntimeError("disque plein, pour le test")

        repet_export.ecrire_repet = echouer
        try:
            resultats = docx_export.executer(self.base)
        finally:
            repet_export.ecrire_repet = origine

        resultat = resultats[0]

        self.assertEqual(resultat.statut, config.STATUT_TERMINE)
        self.assertTrue(self.chemins.docx.exists())
        self.assertFalse(self.chemins.repet.exists())
        self.assertTrue(
            any("répétition" in a for a in resultat.avertissements),
            resultat.avertissements,
        )

    def test_une_anomalie_de_structure_remonte_au_rapport(self):
        """
        Un texte sans personnage annoncé est signalé par `repet_export`. Cet
        avertissement doit atteindre le rapport de l'étape : confiné au JSON, il
        ne serait jamais lu.
        """
        io.ecrire_texte_atomique(
            self.chemins.edit,
            "**ACTE PREMIER**\nUne phrase orpheline.\n\n**JAN.**\nBonjour.\n",
        )

        resultats = docx_export.executer(self.base)

        self.assertEqual(resultats[0].statut, config.STATUT_TERMINE)
        self.assertTrue(
            any("sans personnage" in a for a in resultats[0].avertissements),
            resultats[0].avertissements,
        )


# ============================================================
# 6. CAS PARTICULIERS ET ROBUSTESSE
# ============================================================


@unittest.skipUnless(DOCX_DISPONIBLE, "python-docx n'est pas installé")
class TestRobustesse(unittest.TestCase):
    def setUp(self):
        self._verbosite = config.VERBOSITE
        config.VERBOSITE = 0
        self._dossier = tempfile.TemporaryDirectory()
        self.base = Path(self._dossier.name)
        self.chemins = io.resoudre_chemins("Le Malentendu", self.base)

    def tearDown(self):
        self._dossier.cleanup()
        config.VERBOSITE = self._verbosite

    def test_aucun_fichier_edite(self):
        self.assertEqual(docx_export.executer(self.base), [])

    def test_fichier_vide_signale(self):
        io.ecrire_texte_atomique(self.chemins.edit, "   \n\n")

        resultats = docx_export.executer(self.base)

        self.assertEqual(resultats[0].statut, config.STATUT_ECHEC)
        self.assertIn("vide", resultats[0].erreur)
        self.assertFalse(self.chemins.docx.exists())

    def test_marqueur_de_bloc_en_echec_signale(self):
        """
        Le DOCX est tout de même produit — mieux vaut un document incomplet
        signalé qu'aucun document — mais l'anomalie est remontée.
        """
        io.ecrire_texte_atomique(
            self.chemins.edit,
            "**JAN.**\nBonjour.\n\n"
            + config.MARQUEUR_ECHEC_BLOC.format(numero=3)
            + "\n",
        )

        resultats = docx_export.executer(self.base)

        self.assertTrue(self.chemins.docx.exists())
        self.assertTrue(
            any("bloc en échec" in a for a in resultats[0].avertissements)
        )

    def test_classement_incertain_remonte(self):
        io.ecrire_texte_atomique(
            self.chemins.edit,
            "**ACTE PREMIER**\n\n**JAN.**\nA.\n\n**JAN.**\nB.\n\n"
            "**LA VOIX**\n\n*Silence.*\n",
        )

        resultats = docx_export.executer(self.base)

        self.assertEqual(resultats[0].classements_incertains, ["LA VOIX"])

    def test_journal_ecrit(self):
        io.ecrire_texte_atomique(self.chemins.edit, PIECE)

        docx_export.executer(self.base)

        journal = io.lire_sidecar(io.dossier_temporaire(self.base) / "journal_docx.json")

        self.assertEqual(journal["etape"], "docx")
        bilan = journal["livres"]["Le Malentendu"]
        self.assertEqual(bilan["actes"], 2)
        self.assertEqual(journal["configuration"]["police"], config.POLICE_TEXTE)

    def test_style_manquant_leve_une_erreur_explicite(self):
        """
        Contrat avec `blocks` : un type de ligne sans style doit échouer
        immédiatement, pas produire un paragraphe au style par défaut.
        """
        with self.assertRaises(KeyError):
            docx_export.nom_style(blocks.TypeLigne.VIDE)

    def test_alignement_inconnu_leve_une_erreur_explicite(self):
        with self.assertRaises(ValueError):
            docx_export._alignement("diagonal")


if __name__ == "__main__":
    unittest.main(verbosity=2)
