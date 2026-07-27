"""
Étape 4 — Génération DOCX : `<Livre>_EDIT.txt` → `<Livre>.docx`.

**Aucune IA.** Uniquement `python-docx` et la convention typographique établie
à l'étape 2. Deux exécutions sur le même `EDIT.txt` produisent le même
document : c'est ce déterminisme qui permet de régénérer un DOCX après avoir
changé une marge, sans repayer un seul appel.

Le traitement est en **deux temps**, et cet ordre n'est pas négociable :

1. `blocks.construire_index_structure()` parcourt le document entier et décide
   du type de chaque ligne en gras. Le type d'un `**UN.**` est une propriété
   *globale* — acte dans une pièce en parties numérotées, scène dans une pièce
   qui possède déjà des `ACTE I` — il ne peut donc pas se décider ligne à ligne.
2. Chaque ligne est ensuite classée et rendue selon son style.

Le parsing est lui-même à deux niveaux : le **niveau ligne** détermine le style
de paragraphe, le **niveau run** traite les emphases internes (`decouper_en_runs`).
Sans le second, une réplique du type « Je t'attendais *elle se lève* depuis une
heure » perdrait sa didascalie intercalée.

La table d'inspection est affichée **avant** la génération, afin que vous voyiez
ce que le parseur a compris — notamment les classements incertains — plutôt que
de le découvrir à la première page blanche parasite.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from functools import lru_cache
from pathlib import Path
from typing import Any

from theatre_editor import config
from theatre_editor.utils import blocks, io
from theatre_editor.utils import logging as journalisation

NOM_ETAPE = "docx"

# Types de ligne sans rendu propre : ils ne produisent aucun paragraphe.
TYPES_IGNORES = frozenset({blocks.TypeLigne.VIDE})

# Texte du séparateur de scène, rendu comme une didascalie centrée.
TEXTE_SEPARATEUR = "*"


# ============================================================
# 1. RÉSULTATS
# ============================================================


@dataclass
class ResultatLivre:
    """Bilan de la génération d'un document."""

    nom: str
    statut: str = config.STATUT_TERMINE
    paragraphes: int = 0
    actes: int = 0
    scenes: int = 0
    personnages: int = 0
    classements_incertains: list[str] = field(default_factory=list)
    avertissements: list[str] = field(default_factory=list)
    duree_secondes: float = 0.0
    erreur: str | None = None

    def champs_journal(self) -> dict[str, Any]:
        return {
            "statut": self.statut,
            "paragraphes": self.paragraphes,
            "actes": self.actes,
            "scenes": self.scenes,
            "personnages": self.personnages,
            "classements_incertains": self.classements_incertains,
            "avertissements": self.avertissements,
            "duree_secondes": self.duree_secondes,
            "erreur": self.erreur,
        }


# ============================================================
# 2. DÉPENDANCE python-docx
# ============================================================


@lru_cache(maxsize=1)
def _module_docx():
    """
    Importe `python-docx` à la demande.

    Import différé pour la même raison que dans les autres modules : permettre
    de charger `theatre_editor` sans que toutes les dépendances soient
    installées.
    """
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.shared import Cm, Pt
        from docx.oxml.ns import qn

        return {
            "docx": docx,
            "ALIGN": WD_ALIGN_PARAGRAPH,
            "STYLE_TYPE": WD_STYLE_TYPE,
            "Cm": Cm,
            "Pt": Pt,
            "qn": qn,
        }
    except ImportError as erreur:  # pragma: no cover - dépend de l'environnement
        raise RuntimeError(
            "python-docx est introuvable.\n"
            "Installez-le :  pip install -U python-docx"
        ) from erreur


def _alignement(nom: str):
    """Traduit un alignement de `config` en énumération python-docx."""
    outils = _module_docx()

    correspondances = {
        "centre": outils["ALIGN"].CENTER,
        "justifie": outils["ALIGN"].JUSTIFY,
        "gauche": outils["ALIGN"].LEFT,
    }

    if nom not in correspondances:
        raise ValueError(
            f"alignement inconnu dans config.DEFINITIONS_STYLES : « {nom} ». "
            f"Valeurs admises : {', '.join(sorted(correspondances))}."
        )

    return correspondances[nom]


# ============================================================
# 3. CONSTRUCTION DU DOCUMENT
# ============================================================


def creer_document():
    """Crée un document vierge, avec ses marges et ses styles."""
    outils = _module_docx()

    document = outils["docx"].Document()

    _appliquer_marges(document)
    _creer_styles(document)

    return document


def _appliquer_marges(document) -> None:
    """
    Applique des marges généreuses sur les quatre côtés.

    Aucun numéro de page n'est ajouté : `python-docx` n'en insère aucun par
    défaut, l'exigence est donc satisfaite par abstention — il n'y a rien à
    faire, seulement rien à ajouter.
    """
    centimetres = _module_docx()["Cm"](config.MARGE_CM)

    for section in document.sections:
        section.top_margin = centimetres
        section.bottom_margin = centimetres
        section.left_margin = centimetres
        section.right_margin = centimetres


def _creer_styles(document) -> None:
    """
    Crée les styles de paragraphe décrits dans `config.DEFINITIONS_STYLES`.

    Tout passe par des styles nommés, jamais par une mise en forme appliquée
    run par run : changer le corps des titres de tout un livre reste ainsi la
    modification d'une seule ligne de `config.py`.
    """
    for cle, definition in config.DEFINITIONS_STYLES.items():
        _creer_style(document, cle, definition)


def _creer_style(document, cle: str, definition: dict[str, Any]) -> None:
    """Crée un style unique et le configure entièrement."""
    outils = _module_docx()

    nom = f"{config.PREFIXE_STYLE}{definition['nom']}"

    style = document.styles.add_style(nom, outils["STYLE_TYPE"].PARAGRAPH)
    style.quick_style = True

    police = style.font
    police.name = config.POLICE_TEXTE
    police.size = outils["Pt"](definition["taille_pt"])
    police.bold = bool(definition["gras"])
    police.italic = bool(definition["italique"])
    # Aucune couleur n'est définie : la valeur héritée est le noir automatique.

    _forcer_police(style, config.POLICE_TEXTE)

    format_paragraphe = style.paragraph_format
    format_paragraphe.alignment = _alignement(str(definition["alignement"]))
    format_paragraphe.space_before = outils["Pt"](definition["espace_avant_pt"])
    format_paragraphe.space_after = outils["Pt"](definition["espace_apres_pt"])
    format_paragraphe.page_break_before = bool(definition["saut_de_page"])

    # Évite qu'un nom de personnage se retrouve seul en bas de page, séparé de
    # la réplique qu'il annonce.
    format_paragraphe.keep_with_next = cle in {
        "titre_acte",
        "titre_scene",
        "distribution",
        "personnage",
    }


def _forcer_police(style, nom_police: str) -> None:
    """
    Renseigne les variantes `eastAsia` et `cs` de la police.

    `python-docx` ne remplit que l'attribut `ascii` de `w:rFonts`. Word peut
    alors substituer une autre police sur certains caractères — typiquement les
    guillemets ou les tirets cadratins d'un texte français. Ce détail est
    invisible à la génération et très visible à l'impression.
    """
    qualifier = _module_docx()["qn"]
    element = style.element.rPr.rFonts

    for attribut in ("w:eastAsia", "w:cs", "w:hAnsi"):
        element.set(qualifier(attribut), nom_police)


# ============================================================
# 4. RENDU DES LIGNES
# ============================================================


def nom_style(type_ligne: blocks.TypeLigne) -> str:
    """
    Retourne le nom du style associé à un type de ligne.

    Les valeurs de `TypeLigne` correspondant exactement aux clés de
    `DEFINITIONS_STYLES`, un type sans style provoque une erreur immédiate
    plutôt qu'un paragraphe au style par défaut passé inaperçu.
    """
    definition = config.DEFINITIONS_STYLES.get(type_ligne.value)

    if definition is None:
        raise KeyError(
            f"aucun style défini pour le type « {type_ligne.value} » : "
            "complétez config.DEFINITIONS_STYLES."
        )

    return f"{config.PREFIXE_STYLE}{definition['nom']}"


def ajouter_paragraphe(document, ligne: blocks.LigneClassee) -> bool:
    """
    Ajoute un paragraphe correspondant à une ligne classée.

    Returns:
        True si un paragraphe a été créé.
    """
    if ligne.type in TYPES_IGNORES:
        return False

    if ligne.type is blocks.TypeLigne.SEPARATEUR:
        paragraphe = document.add_paragraph(style=nom_style(blocks.TypeLigne.DIDASCALIE))
        paragraphe.add_run(TEXTE_SEPARATEUR)
        return True

    paragraphe = document.add_paragraph(style=nom_style(ligne.type))

    _ajouter_runs(paragraphe, ligne)

    return True


def _ajouter_runs(paragraphe, ligne: blocks.LigneClassee) -> None:
    """
    Ajoute le contenu d'une ligne, en respectant ses emphases internes.

    Les emphases ne sont découpées que dans le corps du texte. Sur un titre ou
    un nom de personnage, le style porte déjà le gras, et les marqueurs ont été
    retirés par `contenu_sans_marqueurs()` : il n'y a rien à interpréter.
    """
    if ligne.type is not blocks.TypeLigne.TEXTE:
        paragraphe.add_run(ligne.texte)
        return

    for fragment in blocks.decouper_en_runs(ligne.texte):
        run = paragraphe.add_run(fragment.texte)

        # On ne force l'attribut que lorsqu'il est demandé : laisser None
        # préserve la valeur héritée du style.
        if fragment.gras:
            run.bold = True

        if fragment.italique:
            run.italic = True


# ============================================================
# 5. GÉNÉRATION D'UN LIVRE
# ============================================================


def construire_docx(texte: str) -> tuple[Any, blocks.IndexStructure, int]:
    """
    Construit le document à partir d'un texte édité.

    Returns:
        `(document, index de structure, nombre de paragraphes)`.
    """
    index = blocks.construire_index_structure(texte)
    document = creer_document()

    paragraphes = 0

    for ligne in blocks.classifier_document(texte, index):
        if ajouter_paragraphe(document, ligne):
            paragraphes += 1

    return document, index, paragraphes


def traiter_livre(
    chemin_edit: Path,
    journal: journalisation.Journal,
) -> ResultatLivre:
    """Génère le DOCX d'un livre à partir de son fichier édité."""
    nom_livre = io.nom_livre_depuis_edit(chemin_edit)
    chemins = io.resoudre_chemins(nom_livre, chemin_edit.parent)
    resultat = ResultatLivre(nom=nom_livre)

    journalisation.section(f"DOCX — {nom_livre}")

    with journalisation.Chrono() as chrono:
        try:
            _generer(chemins=chemins, resultat=resultat)
        except Exception as erreur:
            resultat.statut = config.STATUT_ECHEC
            resultat.erreur = str(erreur)
            journalisation.echec(f"{nom_livre} : {erreur}")

    resultat.duree_secondes = chrono.secondes

    journal.resumer_livre(nom_livre, **resultat.champs_journal())
    journal.sauvegarder()

    return resultat


def _generer(*, chemins: io.CheminsLivre, resultat: ResultatLivre) -> None:
    """Corps de la génération d'un livre."""
    texte = io.lire_texte(chemins.edit)

    if not texte.strip():
        raise ValueError(f"{chemins.edit.name} est vide")

    document, index, paragraphes = construire_docx(texte)

    # La table d'inspection est affichée AVANT l'écriture : vous voyez ce que le
    # parseur a compris au moment où le document peut encore être rejeté.
    journalisation.info("")
    journalisation.info(blocks.rapport_classification(index))
    journalisation.info("")

    document.save(str(chemins.docx))

    resultat.paragraphes = paragraphes
    resultat.actes = index.compter(blocks.TypeLigne.TITRE_ACTE)
    resultat.scenes = index.compter(blocks.TypeLigne.TITRE_SCENE)
    resultat.personnages = index.compter(blocks.TypeLigne.PERSONNAGE)
    resultat.classements_incertains = [c.affichage for c in index.incertains]
    resultat.avertissements = list(index.avertissements)

    if config.MARQUEUR_ECHEC_BLOC.split("{")[0] in texte:
        resultat.avertissements.append(
            "le texte édité contient un marqueur de bloc en échec : "
            "relancez l'étape « edition »"
        )

    for avertissement in resultat.avertissements:
        journalisation.alerte(avertissement)

    journalisation.succes(
        f"{chemins.docx.name} — {paragraphes} paragraphes, "
        f"{resultat.actes} acte(s), {resultat.scenes} scène(s), "
        f"{resultat.personnages} personnage(s)"
    )


# ============================================================
# 6. POINT D'ENTRÉE DE L'ÉTAPE
# ============================================================


def executer(dossier: Path | None = None) -> list[ResultatLivre]:
    """
    Génère les DOCX de tous les livres édités du dossier.

    Args:
        dossier: dossier à parcourir. `config.DOSSIER_DRIVE` par défaut.

    Returns:
        Un bilan par livre traité.
    """
    base = dossier if dossier is not None else config.DOSSIER_DRIVE

    journalisation.titre("Étape 4 — Génération DOCX")

    fichiers = io.lister_fichiers_edit(base)
    journalisation.info(f"Dossier : {base}")
    journalisation.info(f"Fichiers édités trouvés : {len(fichiers)}")

    if not fichiers:
        journalisation.alerte(
            f"aucun fichier « {config.SUFFIXE_EDIT} » — "
            "lancez d'abord l'étape « edition »"
        )
        return []

    journal = journalisation.Journal.charger_ou_creer(
        NOM_ETAPE,
        base,
        {
            "police": config.POLICE_TEXTE,
            "taille_texte_pt": config.TAILLE_TEXTE_PT,
            "taille_titre_acte_pt": config.TAILLE_TITRE_ACTE_PT,
            "taille_titre_scene_pt": config.TAILLE_TITRE_SCENE_PT,
            "marge_cm": config.MARGE_CM,
            "saut_de_page_avant_acte": config.SAUT_DE_PAGE_AVANT_ACTE,
        },
    )

    resultats = [traiter_livre(chemin, journal) for chemin in fichiers]

    _afficher_recapitulatif(resultats, journal)

    return resultats


def _afficher_recapitulatif(
    resultats: list[ResultatLivre],
    journal: journalisation.Journal,
) -> None:
    """Affiche le bilan global de l'étape."""
    journalisation.recapitulatif(
        {
            "Documents générés": sum(
                1 for r in resultats if r.statut == config.STATUT_TERMINE
            ),
            "Échecs": sum(1 for r in resultats if r.statut == config.STATUT_ECHEC),
            "Paragraphes": sum(r.paragraphes for r in resultats),
            "Actes": sum(r.actes for r in resultats),
            "Scènes": sum(r.scenes for r in resultats),
            "Personnages": sum(r.personnages for r in resultats),
            "Durée": journalisation.formater_duree(
                sum(r.duree_secondes for r in resultats)
            ),
            "Journal": journal.chemin.name,
        }
    )

    incertains = [r.nom for r in resultats if r.classements_incertains]

    if incertains:
        journalisation.info("")
        journalisation.alerte(
            "classements incertains à relire dans la table d'inspection : "
            f"{', '.join(incertains)}"
        )
