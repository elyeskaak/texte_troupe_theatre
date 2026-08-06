"""
Superpose, en trois couleurs, les coupes d'un relecteur et les coupes calculées.

Sert à comparer côte à côte, dans un seul `.docx` à suivi de modifications :

- les coupes d'un **relecteur humain**, déjà marquées dans un `.docx` annoté
  (suppressions Word `<w:del>` sous son nom) → **vert** ;
- mes coupes **passe 1** (structurelles) → **rouge** ;
- mes coupes **passe 2** (rabotage fin) → **bleu**.

Quand un même passage est coupé par le relecteur ET par moi, il est attribué au
**relecteur** (priorité) : mes couleurs ne marquent donc que *mes ajouts*, ce
qu'il n'avait pas coupé. On voit ainsi d'un coup d'œil l'apport de chaque source.

Usage :

    python fusion_coupes.py <source.docx> <relecteur.docx> <coupes.json> --sortie <fichier.docx>

`source.docx` : le texte mis en forme d'origine (base du rendu, format préservé).
`relecteur.docx` : le même texte annoté de suppressions en suivi de modifications.
`coupes.json` : mes coupes (schéma de `appliquer_coupes`).
"""

from __future__ import annotations

import argparse
import difflib
import json
import sys
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn

import appliquer_coupes as ac

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

# Catégorie du relecteur, en plus de mes passes (1 rouge, 2 bleu). On ne mute
# pas les tables globales de appliquer_coupes : on passe ces tables étendues à
# appliquer_sur_run, ce qui laisse l'outil de base intact.
CAT_RELECTEUR = 3
COULEURS = {**ac.COULEURS, CAT_RELECTEUR: "008000"}  # vert pour le relecteur
AUTEURS = {**ac.AUTEURS, CAT_RELECTEUR: "Elyes Kaak (relecture)"}


def masque_relecteur(chemin: str) -> tuple[str, list[bool]]:
    """
    Reconstruit le texte plat d'un `.docx` annoté et le masque de ses coupes.

    Même convention que `appliquer_coupes.lire_document` (runs concaténés,
    paragraphes séparés par un saut de ligne) pour que les offsets coïncident.
    Les runs enveloppés dans un `<w:del>` sont marqués coupés ; les `<w:ins>`
    (ajouts du relecteur) sont ignorés, pour que le texte reste celui d'origine.
    """
    document = Document(chemin)
    parts: list[str] = []
    coupe: list[bool] = []

    for paragraphe in document.paragraphs:
        for enfant in paragraphe._p.iterchildren():
            if enfant.tag == qn("w:ins"):
                continue
            dans_del = enfant.tag == qn("w:del")
            for run in enfant.iter(qn("w:r")):
                for texte in run.iter():
                    if texte.tag in (qn("w:t"), qn("w:delText")) and texte.text:
                        parts.append(texte.text)
                        coupe.extend([dans_del] * len(texte.text))
        parts.append("\n")
        coupe.append(False)

    return "".join(parts), coupe


def aligner_masque(texte_source: str, texte_relecteur: str, coupe_relecteur: list[bool]) -> list[bool]:
    """
    Reporte le masque de coupes du relecteur sur les offsets du texte source.

    Le fichier relecteur n'est pas rigoureusement identique à la source (le
    relecteur y a fait quelques insertions et corrections). On aligne les deux
    textes — normalisés pour ignorer les variantes d'apostrophe et d'espace — et
    on ne reporte les coupes que sur les tranches communes aux deux. Les rares
    zones divergentes (les corrections du relecteur) ne portent pas de coupe.
    """
    # Alignement ligne par ligne : sur ~4 000 lignes variées difflib est
    # instantané, là où un diff caractère par caractère sur 140 000 caractères
    # d'un alphabet minuscule serait quadratique. Les lignes communes aux deux
    # versions ont un contenu identique — donc la même longueur — et se reportent
    # caractère pour caractère ; les lignes divergentes (corrections du relecteur)
    # ne portent pas de coupe.
    # On découpe sur les vrais sauts de ligne AVANT de normaliser (la
    # normalisation transforme justement « \n » en espace). Chaque ligne est
    # ensuite normalisée — apostrophes, espaces insécables — pour que deux lignes
    # au même contenu s'apparient malgré ces variantes. La normalisation
    # conservant les longueurs, les offsets restent ceux du texte d'origine.
    brut_a, brut_b = texte_relecteur.split("\n"), texte_source.split("\n")
    lignes_a = [ac._normaliser(ligne) for ligne in brut_a]
    lignes_b = [ac._normaliser(ligne) for ligne in brut_b]

    def debuts(lignes):
        position, sortie = 0, []
        for ligne in lignes:
            sortie.append(position)
            position += len(ligne) + 1
        return sortie

    debut_a, debut_b = debuts(brut_a), debuts(brut_b)
    norm_a = "\n".join(lignes_a)
    norm_b = "\n".join(lignes_b)
    coupe = [False] * len(texte_source)

    def reporter_caractere(a0, a1, b0, b1):
        """Aligne finement deux tranches (peu de caractères) et reporte les coupes."""
        for tag, i1, i2, j1, j2 in difflib.SequenceMatcher(
            None, norm_a[a0:a1], norm_b[b0:b1], autojunk=False
        ).get_opcodes():
            if tag == "equal":
                for k in range(i2 - i1):
                    if coupe_relecteur[a0 + i1 + k]:
                        coupe[b0 + j1 + k] = True

    for tag, i1, i2, j1, j2 in difflib.SequenceMatcher(None, lignes_a, lignes_b, autojunk=False).get_opcodes():
        if tag == "equal":
            for k in range(i2 - i1):
                base_a, base_b = debut_a[i1 + k], debut_b[j1 + k]
                for x in range(len(lignes_a[i1 + k])):
                    if coupe_relecteur[base_a + x]:
                        coupe[base_b + x] = True
        elif tag == "replace":
            # lignes retouchées par le relecteur : on descend au caractère, sur
            # cette tranche seulement (courte, donc pas de coût quadratique global).
            a0 = debut_a[i1]
            a1 = debut_a[i2] - 1 if i2 < len(debut_a) else len(norm_a)
            b0 = debut_b[j1]
            b1 = debut_b[j2] - 1 if j2 < len(debut_b) else len(norm_b)
            reporter_caractere(a0, a1, b0, b1)
    return coupe


def categories_fusionnees(texte: str, mes_plages, coupe_relecteur: list[bool]) -> list[ac.Plage]:
    """
    Fond mes plages et le masque du relecteur en plages non chevauchantes.

    Priorité au relecteur : un caractère qu'il a coupé prend sa catégorie, quelle
    que soit la mienne. Ailleurs, ma passe (1 ou 2) s'applique.
    """
    categorie = [0] * len(texte)
    for plage in mes_plages:
        for i in range(plage.debut, plage.fin):
            categorie[i] = plage.passe
    for i, coupe in enumerate(coupe_relecteur):
        if coupe:
            categorie[i] = CAT_RELECTEUR

    plages: list[ac.Plage] = []
    i = 0
    while i < len(categorie):
        if categorie[i]:
            j = i
            while j < len(categorie) and categorie[j] == categorie[i]:
                j += 1
            plages.append(ac.Plage(i, j, categorie[i]))
            i = j
        else:
            i += 1
    return plages


def main(argv: list[str] | None = None) -> int:
    analyseur = argparse.ArgumentParser(description="Fusionne coupes relecteur + coupes calculées en 3 couleurs.")
    analyseur.add_argument("source")
    analyseur.add_argument("relecteur")
    analyseur.add_argument("coupes")
    analyseur.add_argument("--sortie", required=True)
    options = analyseur.parse_args(argv)

    document = Document(options.source)
    texte, runs, paras = ac.lire_document(document)

    texte_relecteur, coupe_brut = masque_relecteur(options.relecteur)
    coupe_relecteur = aligner_masque(texte, texte_relecteur, coupe_brut)

    with open(options.coupes, encoding="utf-8") as f:
        coupes = json.load(f)["coupes"]
    mes_plages = ac.propager_noms(paras, ac.resoudre_plages(texte, coupes))

    plages = categories_fusionnees(texte, mes_plages, coupe_relecteur)

    date_iso = datetime.now().isoformat(timespec="seconds")
    for info in runs:
        ac.appliquer_sur_run(info, plages, date_iso, couleurs=COULEURS, auteurs=AUTEURS)
    document.save(options.sortie)

    import re
    def mots(cat):
        return sum(len(re.findall(r"\S+", texte[p.debut:p.fin])) for p in plages if p.passe == cat)
    print(f"{options.sortie} écrit.")
    print(f"Elyes (vert) : {mots(3)} mots | moi passe 1 (rouge) : {mots(1)} | passe 2 (bleu) : {mots(2)}")
    print(f"Total barré : {mots(1) + mots(2) + mots(3)} mots.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
