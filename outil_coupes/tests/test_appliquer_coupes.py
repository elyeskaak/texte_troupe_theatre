"""
Tests de `appliquer_coupes`.

La résolution des plages et la propagation des noms sont testées sans dépendance.
L'application sur un vrai `.docx` exige `python-docx` (dépendance réelle de
l'outil) : ce test se saute proprement si le module est absent.
"""
from __future__ import annotations

import unittest

import appliquer_coupes as ac


class ResoudrePlages(unittest.TestCase):
    TEXTE = "JAN.\nBonjour, comment vas-tu ce matin.\nMARTHA.\nTrès bien.\n"

    def test_texte_exact(self):
        plages = ac.resoudre_plages(self.TEXTE, [{"passe": 1, "texte": ", comment vas-tu"}])
        self.assertEqual(self.TEXTE[plages[0].debut:plages[0].fin], ", comment vas-tu")
        self.assertEqual(plages[0].passe, 1)

    def test_debut_fin_inclut_la_fin(self):
        plages = ac.resoudre_plages(self.TEXTE, [{"passe": 2, "debut": "comment", "fin": "ce matin"}])
        self.assertEqual(self.TEXTE[plages[0].debut:plages[0].fin], "comment vas-tu ce matin")

    def test_ancre_introuvable_leve(self):
        with self.assertRaises(ValueError):
            ac.resoudre_plages(self.TEXTE, [{"passe": 1, "texte": "absent du texte"}])

    def test_ancre_ambigue_leve(self):
        with self.assertRaises(ValueError):
            ac.resoudre_plages("Oui. Oui. Oui.", [{"passe": 1, "texte": "Oui."}])

    def test_passe_inconnue_leve(self):
        with self.assertRaises(ValueError):
            ac.resoudre_plages(self.TEXTE, [{"passe": 3, "texte": "Bonjour"}])

    def test_chevauchement_leve(self):
        coupes = [{"passe": 1, "texte": "Bonjour, comment"}, {"passe": 2, "texte": "comment vas-tu"}]
        with self.assertRaises(ValueError):
            ac.resoudre_plages(self.TEXTE, coupes)

    def test_ancre_droite_matche_apostrophe_typographique(self):
        texte = "Je l’ai vu hier soir."  # apostrophe typographique dans le document
        plages = ac.resoudre_plages(texte, [{"passe": 1, "texte": "l'ai vu"}])  # ancre droite
        self.assertEqual(texte[plages[0].debut:plages[0].fin], "l’ai vu")


class PasseAPosition(unittest.TestCase):
    def test_dedans_et_dehors(self):
        plages = [ac.Plage(2, 5, 1)]
        self.assertEqual(ac.passe_a_position(1, plages), 0)
        self.assertEqual(ac.passe_a_position(2, plages), 1)
        self.assertEqual(ac.passe_a_position(4, plages), 1)
        self.assertEqual(ac.passe_a_position(5, plages), 0)


class PropagerNoms(unittest.TestCase):
    # Offsets cohérents avec un texte plat « JAN.\nBonjour.\nMARTHA.\nSalut.\n ».
    def _paras(self):
        return [
            ac.InfoParagraphe(0, "JAN.", True, False),
            ac.InfoParagraphe(5, "Bonjour.", False, False),
            ac.InfoParagraphe(14, "MARTHA.", True, False),
            ac.InfoParagraphe(22, "Salut.", False, False),
        ]

    def test_nom_coupe_si_replique_entierement_coupee(self):
        plages = [ac.Plage(5, 13, 1)]  # « Bonjour. »
        etendu = ac.propager_noms(self._paras(), plages)
        self.assertTrue(any(p.debut == 0 and p.fin == 4 for p in etendu))  # JAN. emporté
        self.assertFalse(any(p.debut == 14 for p in etendu))  # MARTHA. conservé

    def test_nom_conserve_si_replique_partiellement_coupee(self):
        plages = [ac.Plage(5, 9, 1)]  # « Bonj » seulement
        etendu = ac.propager_noms(self._paras(), plages)
        self.assertFalse(any(p.debut == 0 for p in etendu))

    def test_titre_non_emporte_par_une_didascalie(self):
        paras = [
            ac.InfoParagraphe(0, "ACTE I.", True, False),
            ac.InfoParagraphe(8, "Une rue.", False, True),  # didascalie : italique
        ]
        plages = [ac.Plage(8, 16, 1)]
        etendu = ac.propager_noms(paras, plages)
        self.assertFalse(any(p.debut == 0 for p in etendu))


try:
    import docx  # noqa: F401

    _DOCX = True
except ImportError:
    _DOCX = False


@unittest.skipUnless(_DOCX, "python-docx non installé")
class ApplicationSurDocx(unittest.TestCase):
    def _docx_source(self, tmp):
        from docx import Document
        from docx.shared import Pt

        document = Document()
        nom = document.add_paragraph().add_run("MARTHA.")
        nom.bold = True
        nom.font.name = "EB Garamond"
        nom.font.size = Pt(15)
        replique = document.add_paragraph().add_run("Bonjour, vraiment tout le monde.")
        replique.font.name = "EB Garamond"
        replique.font.size = Pt(15)

        from pathlib import Path
        chemin = Path(tmp) / "src.docx"
        document.save(str(chemin))
        return str(chemin)

    def test_coupe_partielle_preserve_le_format_et_le_nom(self):
        import tempfile
        import zipfile

        with tempfile.TemporaryDirectory() as tmp:
            source = self._docx_source(tmp)
            document, texte, plages = ac.construire_document(source, [{"passe": 1, "texte": ", vraiment"}])
            from pathlib import Path
            out = Path(tmp) / "out.docx"
            document.save(str(out))
            xml = zipfile.ZipFile(out).read("word/document.xml").decode("utf-8")

        self.assertIn("<w:del ", xml)
        self.assertIn("Coupe passe 1", xml)
        self.assertIn("C00000", xml)
        self.assertIn("EB Garamond", xml)  # format d'origine préservé
        # coupe partielle : le nom MARTHA. n'est pas emporté
        self.assertEqual(xml.count("Coupe passe 1"), 1)

    def test_replique_entierement_coupee_emporte_le_nom(self):
        import tempfile
        import zipfile

        with tempfile.TemporaryDirectory() as tmp:
            source = self._docx_source(tmp)
            document, texte, plages = ac.construire_document(
                source, [{"passe": 2, "texte": "Bonjour, vraiment tout le monde."}]
            )
            from pathlib import Path
            out = Path(tmp) / "out.docx"
            document.save(str(out))
            xml = zipfile.ZipFile(out).read("word/document.xml").decode("utf-8")

        # le nom MARTHA. est emporté : deux suppressions (nom + réplique), en bleu
        self.assertIn("Coupe passe 2", xml)
        self.assertIn("0070C0", xml)
        self.assertGreaterEqual(xml.count("<w:del "), 2)


if __name__ == "__main__":
    unittest.main()
