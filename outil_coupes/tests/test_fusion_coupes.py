"""
Tests de `fusion_coupes` : alignement du masque relecteur, priorité, relecture
des suppressions d'un .docx annoté.
"""
from __future__ import annotations

import unittest

import appliquer_coupes as ac
import fusion_coupes as fc


class AlignerMasque(unittest.TestCase):
    def test_report_malgre_une_ligne_inseree(self):
        source = "Bonjour\nComment vas-tu\nAu revoir"
        # le relecteur a inséré une ligne : les offsets diffèrent
        relecteur = "Bonjour\nSalut, ligne en plus\nComment vas-tu\nAu revoir"
        coupe = [False] * len(relecteur)
        deb = relecteur.find("Comment vas-tu")
        for k in range(len("Comment vas-tu")):
            coupe[deb + k] = True

        masque = fc.aligner_masque(source, relecteur, coupe)

        j = source.find("Comment vas-tu")
        self.assertTrue(all(masque[j + k] for k in range(len("Comment vas-tu"))))
        self.assertFalse(masque[source.find("Bonjour")])

    def test_apostrophe_typographique_ne_casse_pas_l_alignement(self):
        source = "Je l'ai vu\nFin"
        relecteur = "Je l’ai vu\nFin"  # apostrophe typographique
        coupe = [c != "\n" for c in relecteur[: relecteur.find("\n")]] + [False] * (
            len(relecteur) - relecteur.find("\n")
        )
        masque = fc.aligner_masque(source, relecteur, coupe)
        self.assertTrue(masque[0])  # « Je... » reporté


class CategoriesFusionnees(unittest.TestCase):
    def test_priorite_au_relecteur(self):
        texte = "AAAA BBBB CCCC"  # A:0-3  espace:4  B:5-8  espace:9  C:10-13
        mes = [ac.Plage(0, 4, 1), ac.Plage(10, 14, 2)]
        coupe_rel = [i < 9 for i in range(len(texte))]  # coupe « AAAA BBBB »

        plages = fc.categories_fusionnees(texte, mes, coupe_rel)
        par_bornes = {(p.debut, p.fin): p.passe for p in plages}

        # « AAAA BBBB » revient au relecteur (3), même sur ma passe 1
        self.assertEqual(par_bornes[(0, 9)], fc.CAT_RELECTEUR)
        # ma passe 2, non touchée par le relecteur, reste
        self.assertEqual(par_bornes[(10, 14)], 2)


class MasqueRelecteur(unittest.TestCase):
    def test_relit_les_suppressions_dun_docx_annote(self):
        import tempfile
        from pathlib import Path

        from docx import Document
        from docx.shared import Pt

        with tempfile.TemporaryDirectory() as tmp:
            document = Document()
            run = document.add_paragraph().add_run("Bonjour tout le monde")
            run.font.name = "EB Garamond"
            run.font.size = Pt(15)
            source = Path(tmp) / "src.docx"
            document.save(str(source))

            # produire un docx annoté (suppressions Word) via appliquer_coupes
            annote_doc, _, _ = ac.construire_document(str(source), [{"passe": 1, "texte": "tout le monde"}])
            annote = Path(tmp) / "annote.docx"
            annote_doc.save(str(annote))

            texte, coupe = fc.masque_relecteur(str(annote))

        i = texte.find("tout le monde")
        self.assertTrue(all(coupe[i + k] for k in range(len("tout le monde"))))
        self.assertFalse(coupe[texte.find("Bonjour")])  # « Bonjour » conservé


if __name__ == "__main__":
    unittest.main()
