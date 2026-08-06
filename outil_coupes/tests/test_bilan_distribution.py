"""
Tests de `bilan_distribution` : canonisation (avec alias), comptage du dialogue
restant par personnage, présence par scène, détection des conflits.

La génération d'un petit `.docx` formaté exige `python-docx` : ces tests se
sautent proprement s'il est absent.
"""
from __future__ import annotations

import io
import unittest
from contextlib import redirect_stdout

import bilan_distribution as bd


class Canoniser(unittest.TestCase):
    def test_retire_jeu_et_ponctuation(self):
        self.assertEqual(bd._canoniser("BÉNÉDICT [à part].", {}), "BÉNÉDICT")

    def test_applique_l_alias(self):
        self.assertEqual(bd._canoniser("UN SEIGNEUR", {"UN SEIGNEUR": "ANTONIO"}), "ANTONIO")


try:
    import docx  # noqa: F401

    _DOCX = True
except ImportError:
    _DOCX = False


def _document():
    from docx import Document

    d = Document()

    def gras(t):
        r = d.add_paragraph().add_run(t)
        r.bold = True

    def dit(t):
        d.add_paragraph().add_run(t)

    gras("ACTE I")
    gras("Scène 1")
    gras("JAN.")
    dit("Bonjour tout le monde ici")
    gras("MARTHA.")
    dit("Salut à toi")
    gras("Scène 2")
    gras("JAN.")
    dit("Encore moi")
    return d


@unittest.skipUnless(_DOCX, "python-docx non installé")
class Parcourir(unittest.TestCase):
    def setUp(self):
        self.doc = _document()
        self.texte, _, self.paras = bd.ac.lire_document(self.doc)

    def test_poids_et_presence_sans_coupe(self):
        coupe = [False] * len(self.texte)
        poids, scenes = bd.parcourir(self.doc, coupe, {})
        self.assertEqual(poids["JAN"], 5 + 2)   # « Bonjour tout le monde ici » + « Encore moi »
        self.assertEqual(poids["MARTHA"], 3)    # « Salut à toi »
        # chaque scène porte les mots dits par personnage dans la scène
        self.assertEqual(set(scenes[0][1]), {"JAN", "MARTHA"})
        self.assertEqual(scenes[0][1]["JAN"], 5)
        self.assertEqual(scenes[1][1], {"JAN": 2})

    def test_les_mots_coupes_ne_comptent_plus(self):
        coupe = [False] * len(self.texte)
        debut = self.texte.index("tout le monde")
        for i in range(debut, debut + len("tout le monde")):
            coupe[i] = True
        poids, _ = bd.parcourir(self.doc, coupe, {})
        self.assertEqual(poids["JAN"], 2 + 2)   # « Bonjour ici » + « Encore moi »

    def test_conflit_detecte_quand_un_comedien_joue_deux_roles_dans_une_scene(self):
        coupe = [False] * len(self.texte)
        _, scenes = bd.parcourir(self.doc, coupe, {})
        cast = {"C1": ["JAN", "MARTHA"]}  # même comédien pour les deux → conflit en Scène 1
        sortie = io.StringIO()
        with redirect_stdout(sortie):
            bd.afficher_faisabilite(scenes, cast, {})
        texte = sortie.getvalue()
        self.assertIn("ACTE I Scène 1", texte)
        self.assertIn("JAN + MARTHA", texte)


if __name__ == "__main__":
    unittest.main()
